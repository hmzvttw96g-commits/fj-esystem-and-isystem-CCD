from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(r"C:\金融专硕\论文\硕士毕业论文_Obsidian_MD统一整理_20260512")
OUT_DIR = ROOT / "06_导师沟通材料"
OUT_PATH = OUT_DIR / "导师讨论版_开题框架_20260514.docx"


ACCENT = "1F4E79"
LIGHT_BLUE = "EAF2F8"
LIGHT_YELLOW = "FFF8E5"
LIGHT_GRAY = "F5F7FA"
TEXT = "1F2933"
MUTED = "667085"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=120, start=120, bottom=120, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_border(cell, color="D0D5DD", size="6"):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        element = tc_borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tc_borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_font(run, size=None, bold=None, color=None):
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    if size:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_paragraph(doc, text="", style=None, size=10.5, bold=False, color=TEXT, align=None, space_after=4):
    p = doc.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    set_font(run, size=size, bold=bold, color=color)
    return p


def add_heading(doc, text, level=1):
    style = f"Heading {level}"
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_before = Pt(8 if level == 1 else 4)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    set_font(run, size=15 if level == 1 else 12.5, bold=True, color=ACCENT if level == 1 else TEXT)
    return p


def add_callout(doc, title, body, fill=LIGHT_BLUE):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_border(cell, color="B8CCE4")
    set_cell_margins(cell, 180, 180, 180, 180)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    title_run = p.add_run(title)
    set_font(title_run, size=10.5, bold=True, color=ACCENT)
    p2 = cell.add_paragraph()
    p2.paragraph_format.line_spacing = 1.2
    p2.paragraph_format.space_after = Pt(0)
    body_run = p2.add_run(body)
    set_font(body_run, size=10, color=TEXT)
    add_paragraph(doc, "", space_after=2)


def add_discussion_box(doc, rows=5, prompt="导师意见 / 讨论记录"):
    table = doc.add_table(rows=rows + 1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    header = table.cell(0, 0)
    set_cell_shading(header, LIGHT_YELLOW)
    set_cell_border(header, color="E4C774")
    set_cell_margins(header, 130, 160, 130, 160)
    header.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = header.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(prompt)
    set_font(run, size=10.5, bold=True, color="7A4F01")
    for idx in range(1, rows + 1):
        cell = table.cell(idx, 0)
        set_cell_border(cell, color="E7EAF0")
        set_cell_margins(cell, 160, 160, 160, 160)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(" " * 2)
        set_font(run, size=11, color=TEXT)
    add_paragraph(doc, "", space_after=4)


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    hdr_cells = table.rows[0].cells
    set_repeat_table_header(table.rows[0])
    for idx, header in enumerate(headers):
        cell = hdr_cells[idx]
        set_cell_shading(cell, ACCENT)
        set_cell_border(cell, color=ACCENT)
        set_cell_margins(cell, 120, 120, 120, 120)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(header)
        set_font(run, size=9.5, bold=True, color="FFFFFF")
        if widths:
            cell.width = Cm(widths[idx])
    for row_idx, row in enumerate(rows):
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cell = cells[idx]
            set_cell_shading(cell, "FFFFFF" if row_idx % 2 == 0 else LIGHT_GRAY)
            set_cell_border(cell, color="D0D5DD")
            set_cell_margins(cell, 110, 110, 110, 110)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if len(str(value)) <= 8 else WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(str(value))
            set_font(run, size=9.2, color=TEXT)
            if widths:
                cell.width = Cm(widths[idx])
    add_paragraph(doc, "", space_after=4)
    return table


def add_checklist(doc, items):
    for item in items:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.2)
        p.paragraph_format.space_after = Pt(3)
        r1 = p.add_run("☐ ")
        set_font(r1, size=10.5, color=ACCENT)
        r2 = p.add_run(item)
        set_font(r2, size=10.3, color=TEXT)


def build_document():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()

    section = doc.sections[0]
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2.1)
    section.right_margin = Cm(2.1)

    styles = doc.styles
    styles["Normal"].font.name = "Microsoft YaHei"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    styles["Normal"].font.size = Pt(10.5)

    # Cover
    for _ in range(3):
        add_paragraph(doc, "", space_after=6)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("硕士论文导师讨论稿")
    set_font(r, size=22, bold=True, color=ACCENT)
    p.paragraph_format.space_after = Pt(8)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("开题式框架与问题清单")
    set_font(r, size=15, bold=True, color=TEXT)

    add_paragraph(doc, "", space_after=20)
    meta = [
        ("拟定题目", "金融支持、AI教育供给与人工智能产业发展的转化机制研究——基于福建省的实证分析"),
        ("讨论时间", "2026年5月"),
        ("讨论目的", "确认选题边界、指标口径、数据可行性和实证主线"),
        ("使用方式", "每节右后方/下方留有讨论空间，可现场记录导师意见"),
    ]
    add_table(doc, ["项目", "内容"], meta, widths=[3.2, 12.2])
    add_callout(
        doc,
        "核心提醒",
        "这份材料不是论文定稿，而是为了明天与导师高效讨论。建议优先请导师判断：题目是否需要收窄、AI指标口径是否认可、金融调节效应能否作为主模型。",
        fill=LIGHT_YELLOW,
    )
    add_discussion_box(doc, rows=7, prompt="导师对题目与总体方向的初步意见")

    doc.add_page_break()

    add_heading(doc, "一、研究背景与问题提出", 1)
    add_paragraph(
        doc,
        "人工智能已成为区域产业竞争和产业升级的重要方向。福建省具备电子信息、软件服务、智能制造、数字经济等基础，但与北京、上海、广东、浙江等AI先发地区相比，AI产业集聚、资本支持和成果转化能力仍有提升空间。",
    )
    add_paragraph(
        doc,
        "既有高等教育科类结构与产业结构耦合研究，通常以三次产业分类为基础。本文的难点在于：人工智能不是传统单一产业门类，而是覆盖第一、第二、第三产业的通用技术。因此，论文需要建立更清晰的AI教育供给、AI产业发展和金融支持指标口径。",
    )
    add_table(
        doc,
        ["讨论点", "当前判断", "需导师确认"],
        [
            ("研究对象", "福建省AI领域金融、教育、产业三系统", "是否聚焦福建，还是加入对照省份"),
            ("核心问题", "金融支持是否影响AI教育供给向AI产业发展的转化效率", "是否比“三系统耦合协调”更适合作为主线"),
            ("理论起点", "商润泽教育-产业耦合框架", "是否可作为开题主要理论起点"),
            ("关键难点", "AI指标口径和产业分类边界", "保守/核心/扩展口径是否可接受"),
        ],
        widths=[3.0, 7.0, 5.4],
    )
    add_discussion_box(doc, rows=7, prompt="导师意见：研究背景与问题意识")

    add_heading(doc, "二、拟定题目与研究主线", 1)
    add_callout(
        doc,
        "建议主线",
        "不把论文写成单纯测算D值的“三系统耦合协调研究”，而是围绕“金融支持是否强化AI教育供给向AI产业发展的转化效率”展开。",
        fill=LIGHT_BLUE,
    )
    add_table(
        doc,
        ["方案", "题目", "评价"],
        [
            ("优先", "金融支持、AI教育供给与人工智能产业发展的转化机制研究——基于福建省的实证分析", "金融专硕特色更清楚，机制检验主线明确"),
            ("备选", "金融支持如何影响AI教育供给向AI产业发展的转化——基于福建省及对照省份的实证研究", "问题导向更强，但题目略口语"),
            ("不优先", "福建省人工智能领域金融-教育-产业三系统耦合协调研究", "容易变成单纯测度，金融主线偏弱"),
        ],
        widths=[2.2, 9.0, 4.2],
    )
    add_discussion_box(doc, rows=6, prompt="导师意见：题目是否调整")

    add_heading(doc, "三、研究目标、问题与假设", 1)
    add_table(
        doc,
        ["类型", "内容"],
        [
            ("总体目标", "构建AI领域金融支持、AI教育供给与AI产业发展的分析框架，识别福建省三系统协调短板，并检验金融支持是否提升教育供给向产业发展的转化效率。"),
            ("RQ1", "AI教育是否具有从一般高等教育和计算机教育中相对独立出来的理论与经验依据？"),
            ("RQ2", "福建省AI领域金融支持、AI教育供给和AI产业发展分别处于什么状态？"),
            ("RQ3", "福建省AI领域三系统耦合协调水平如何，短板来自哪个子系统和哪些指标？"),
            ("RQ4", "金融支持是否强化AI教育供给对AI产业发展的促进作用？"),
        ],
        widths=[2.5, 12.9],
    )
    add_table(
        doc,
        ["假设", "表述", "是否保留"],
        [
            ("H1", "AI教育供给水平越高，AI产业发展水平越高。", "建议保留"),
            ("H2", "金融支持水平越高，AI产业发展水平越高。", "建议保留"),
            ("H3", "金融支持能够正向调节AI教育供给对AI产业发展的影响。", "核心假设"),
            ("H4", "福建省协调不足的短板更可能集中在金融资本供给和产业转化环节。", "描述性假设"),
            ("H5", "第三产业AI应用强度越高，AI教育供给向AI产业转化效率越高。", "可作为拓展/异质性"),
        ],
        widths=[2.0, 10.4, 3.0],
    )
    add_discussion_box(doc, rows=7, prompt="导师意见：研究问题与假设")

    doc.add_page_break()

    add_heading(doc, "四、理论框架与机制逻辑", 1)
    add_paragraph(
        doc,
        "本文继承商润泽关于高等教育科类结构与产业结构耦合的理论逻辑，但不直接套用三次产业分类，而是将一般高等教育结构细化为AI教育供给，将一般产业结构细化为AI产业发展，并引入金融支持作为成果转化过程中的关键条件。",
    )
    add_table(
        doc,
        ["理论模块", "核心解释", "可观察证据"],
        [
            ("AI教育独立性", "AI教育在知识体系、课程结构、能力目标和组织形态上具有相对独立性。", "教育部专业备案、高校培养方案、AI课程、硕博点"),
            ("教育到产业", "AI教育供给通过人才、论文、专利、实验室等形成产业可吸收资源。", "AI专业、论文、专利、校企实验室"),
            ("金融转化器", "金融支持通过资本供给、风险分担、创新筛选和治理约束，提升成果转化效率。", "AI融资事件、VC/PE、政府基金、数字金融"),
            ("产业吸纳机制", "第三产业中信息、金融、医疗、教育、政务等部门更易形成AI规模化应用。", "第三产业AI企业/岗位/专利/融资占比"),
        ],
        widths=[3.0, 7.4, 5.0],
    )
    add_callout(
        doc,
        "机制主线",
        "AI教育供给 E -> 人才、知识、专利与科研成果 -> 在金融支持 F 的作用下获得融资、孵化、转化和企业成长机会 -> AI产业发展 I。",
        fill=LIGHT_BLUE,
    )
    add_discussion_box(doc, rows=7, prompt="导师意见：理论框架是否成立")

    add_heading(doc, "五、指标体系初步设计", 1)
    add_paragraph(
        doc,
        "指标体系建议不追求三个系统的机械一一对应，而是采用“教育供给—产业吸纳—金融转化”的功能链条对应。为解决AI边界不清问题，同时设置保守、核心和扩展三类口径做稳健性检验。",
    )
    add_table(
        doc,
        ["共同维度", "教育系统 E", "产业系统 I", "金融系统 F"],
        [
            ("规模基础", "AI相关专业数量、招生规模", "AI企业数量、AI岗位数量", "AI融资事件数、科技金融规模"),
            ("高层次能力", "AI硕博点、高层次师资", "AI高技术企业、专精特新企业", "VC/PE投资、政府引导基金"),
            ("知识与技术产出", "高校AI论文、高校AI专利", "企业AI专利、AI产品/项目", "科技贷款、知识产权质押融资"),
            ("转化连接", "产业学院、校企实验室", "技术合同、应用示范场景", "成果转化基金、概念验证基金"),
            ("应用吸纳", "毕业生去向、AI+X课程", "第三产业AI应用强度、AI岗位需求", "投向AI应用企业的资本支持"),
        ],
        widths=[3.0, 4.2, 4.2, 4.0],
    )
    add_table(
        doc,
        ["系统", "最低可行指标", "数据来源"],
        [
            ("金融 F", "金融业增加值/GDP、存贷款余额/GDP、数字普惠金融指数、AI融资事件数", "统计年鉴、人民银行、北大数字普惠金融指数、企查查/IT桔子/清科"),
            ("教育 E", "AI相关专业数量、AI论文数量、高校AI专利数量、AI相关硕博点数量", "教育部备案、高校官网、CNKI/WoS、国家知识产权局、学位授权点名单"),
            ("产业 I", "AI企业数量、AI企业专利数量、技术合同成交额、AI岗位数量", "企查查/天眼查、专利库、科技统计、招聘平台"),
            ("控制变量", "人均GDP、产业结构、科技财政支出/GDP、数字基础、城镇化率", "统计年鉴、统计公报、财政年鉴、通信管理局"),
        ],
        widths=[2.4, 7.0, 6.0],
    )
    add_discussion_box(doc, rows=8, prompt="导师意见：指标是否认可 / 是否删减")

    doc.add_page_break()
    add_heading(doc, "六、AI口径与三次产业问题", 1)
    add_table(
        doc,
        ["口径", "教育 E", "产业 I", "用途"],
        [
            ("保守口径", "人工智能、智能科学与技术", "主营业务明确包含AI、机器学习、智能算法、大模型等", "最严格定义，避免泛化"),
            ("核心口径", "保守口径 + 数据科学与大数据技术、机器人工程", "名称、经营范围、产品、专利含AI核心关键词", "建议作为基准"),
            ("扩展口径", "核心口径 + 计算机、软件、电子信息、自动化、统计、应用数学", "核心口径 + 智能制造、工业互联网、云计算、数据服务等", "稳健性检验"),
        ],
        widths=[2.4, 4.5, 6.2, 2.3],
    )
    add_callout(
        doc,
        "与商润泽的衔接",
        "商润泽使用三次产业分类研究宏观教育结构匹配；本文研究AI这一通用技术，应采用“官方口径 + 技术关键词 + 应用场景 + 稳健性口径”的复合识别方法。第三产业AI应用强度可以作为产业吸纳机制，而不替代金融主线。",
        fill=LIGHT_YELLOW,
    )
    add_discussion_box(doc, rows=7, prompt="导师意见：AI分类口径与三次产业处理")

    doc.add_page_break()

    add_heading(doc, "七、研究方法与模型方案", 1)
    add_table(
        doc,
        ["方法", "作用", "是否作为主线"],
        [
            ("综合评价指数", "计算F、E、I三个子系统发展水平", "保留"),
            ("耦合协调度模型", "回答三系统协调水平高不高", "保留，但作为描述性测度"),
            ("障碍度模型", "识别短板指标和短板系统", "保留"),
            ("固定效应模型", "检验E、F对I的影响", "保留"),
            ("调节效应模型", "检验金融支持是否强化E -> I", "核心模型"),
            ("空间计量/DID/政策文本", "拓展分析", "暂不作为主线"),
        ],
        widths=[3.8, 8.2, 3.4],
    )
    add_callout(
        doc,
        "核心计量模型",
        "I_it = α + β1 E_it + β2 F_it + β3(E_it × F_it) + Controls_it + μ_i + λ_t + ε_it。若β3为正且显著，说明金融支持提升AI教育供给向AI产业发展的转化效率。",
        fill=LIGHT_BLUE,
    )
    add_table(
        doc,
        ["当前风险", "说明", "讨论请求"],
        [
            ("CCD公式", "三维乘积公式可能存在量表压缩；两两耦合平均法更直观但需谨慎表述", "请导师判断使用哪种公式更稳妥"),
            ("负值和平移", "现有D值对平移常数敏感", "是否改用0-1标准化避免负值"),
            ("熵值法", "不能把偏离度加权写成熵值法", "是否基准熵值法、稳健性等权重"),
            ("黑盒因子", "理论解释价值大，但缺少客观量化数据", "是否降级为政策讨论"),
        ],
        widths=[3.0, 8.4, 4.0],
    )
    add_discussion_box(doc, rows=8, prompt="导师意见：模型与方法")

    add_heading(doc, "八、数据方案与可行性", 1)
    add_table(
        doc,
        ["数据类型", "可得性判断", "处理方案"],
        [
            ("统计年鉴/统计公报", "较高", "用于GDP、产业结构、金融业增加值、技术合同等"),
            ("教育部专业备案/高校官网", "较高但需整理", "用于AI专业数量、硕博点、产业学院"),
            ("CNKI/WoS/专利库", "可得但多需数据库导出", "用于AI论文、高校专利、企业专利"),
            ("企查查/天眼查/IT桔子/清科", "受限或需账号", "用于AI企业、融资事件、VC/PE"),
            ("招聘平台", "可得但偏差较大", "用于AI岗位数量，建议作为增强变量"),
        ],
        widths=[3.4, 4.0, 8.0],
    )
    add_checklist(
        doc,
        [
            "导师是否认可五省市样本：福建、广东、浙江、江苏、上海",
            "导师是否要求扩展到十省样本：增加北京、安徽、山东、湖北、四川",
            "AI产业增加值拿不到时，是否可用AI企业数量、企业专利、岗位数量替代",
            "AI融资金额拿不到时，是否可用AI融资事件数替代",
            "第三产业AI应用强度是否作为增强变量而非主变量",
        ],
    )
    add_discussion_box(doc, rows=8, prompt="导师意见：数据来源与样本范围")

    add_heading(doc, "九、预期创新与可能贡献", 1)
    add_table(
        doc,
        ["贡献类型", "当前表述", "强弱判断"],
        [
            ("研究对象", "将教育-产业耦合研究推进到AI这一战略性新兴技术领域", "较强"),
            ("理论框架", "提出金融支持参与AI教育成果产业化的转化器机制", "较强"),
            ("方法组合", "区分描述性测度与机制检验，避免只测D值", "中等偏强"),
            ("政策价值", "为福建省科技金融、AI产教融合和成果转化政策提供依据", "较强"),
        ],
        widths=[3.0, 9.0, 3.4],
    )
    add_discussion_box(doc, rows=6, prompt="导师意见：创新点是否过强 / 是否需收敛")

    add_heading(doc, "十、明日重点请导师确认的问题", 1)
    add_checklist(
        doc,
        [
            "题目是否采用“金融支持、AI教育供给与人工智能产业发展的转化机制研究”",
            "论文主线是否从“三系统耦合协调”调整为“金融支持调节教育向产业转化”",
            "AI指标是否采用保守、核心、扩展三类口径",
            "第三产业AI应用强度是否只作为拓展机制变量",
            "CCD公式采用三维乘积公式还是两两耦合平均扩展公式",
            "样本选择采用五省市还是十省面板",
            "受限数据库数据是否可以用公开替代变量先行",
        ],
    )
    add_discussion_box(doc, rows=12, prompt="会议记录：导师最终建议与下一步任务")

    section = doc.sections[0]
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer.add_run("硕士论文导师讨论稿｜仅供讨论使用")
    set_font(footer_run, size=8.5, color=MUTED)

    doc.save(OUT_PATH)
    return OUT_PATH


if __name__ == "__main__":
    print(build_document())

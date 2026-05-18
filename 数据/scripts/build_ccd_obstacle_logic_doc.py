from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


THESIS_DIR = Path(r"C:\金融专硕\论文\硕士毕业论文_Obsidian_MD统一整理_20260512")
OUT_DIR = THESIS_DIR / "06_导师沟通材料"
OUT_DOCX = OUT_DIR / "CCD障碍度与机制检验逻辑说明_20260515.docx"

BLUE = "1F4E79"
LIGHT_BLUE = "EAF3FB"
GOLD = "D9A441"
LIGHT_GOLD = "FFF7DF"
GREEN = "2E7D60"
LIGHT_GREEN = "EAF7F0"
RED = "B55454"
LIGHT_RED = "FFF0EF"
GRAY = "6B7280"
BORDER = "D6DEE8"


def set_cn_font(run, size=11, bold=False, color="1F2933"):
    run.font.name = "微软雅黑"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def borders_cell(cell, color=BORDER, size="8"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for side in ["top", "left", "bottom", "right"]:
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), size)
        el.set(qn("w:color"), color)
        borders.append(el)
    tc_pr.append(borders)


def set_cell(cell, text, fill=None, bold=False, color="1F2933", size=10.5, align="left"):
    cell.text = ""
    if fill:
        shade_cell(cell, fill)
    borders_cell(cell)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if align == "center" else WD_ALIGN_PARAGRAPH.LEFT
    for i, line in enumerate(text.split("\n")):
        if i:
            p.add_run().add_break()
        run = p.add_run(line)
        set_cn_font(run, size=size, bold=bold, color=color)


def set_table_widths(table, widths_cm):
    for row in table.rows:
        for cell, width in zip(row.cells, widths_cm):
            cell.width = Cm(width)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10 if level == 1 else 6)
    p.paragraph_format.space_after = Pt(5)
    run = p.add_run(text)
    set_cn_font(run, size=18 if level == 1 else 13.5, bold=True, color=BLUE if level == 1 else "1F2933")
    return p


def add_para(doc, text, size=11, color="1F2933", bold=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.18
    run = p.add_run(text)
    set_cn_font(run, size=size, color=color, bold=bold)
    return p


def add_callout(doc, title, body, fill=LIGHT_GOLD, border=GOLD):
    table = doc.add_table(rows=1, cols=1)
    table.autofit = True
    cell = table.cell(0, 0)
    shade_cell(cell, fill)
    borders_cell(cell, color=border, size="10")
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(title)
    set_cn_font(run, size=12, bold=True, color=BLUE)
    for line in body.split("\n"):
        p = cell.add_paragraph()
        p.paragraph_format.space_after = Pt(1)
        run = p.add_run(line)
        set_cn_font(run, size=10.5)


def add_flow_table(doc):
    rows = [
        ("1", "AI口径界定", "确定什么算AI教育、AI产业和AI金融支持", "测度前提"),
        ("2", "子系统指数", "分别计算 E（教育）、F（金融）、I（产业）发展水平", "基础变量"),
        ("3", "CCD", "回答三系统整体协调水平如何、是否存在协调不足", "描述性测度"),
        ("4", "障碍度", "在指标层面排序：哪些指标或子系统拖累综合发展水平", "短板诊断"),
        ("5", "固定效应", "检验 E、F 是否与 I 的变化存在经验关联", "关联检验"),
        ("6", "调节效应", "检验 F 是否强化 E -> I 的转化关系，核心关注 E × F", "机制检验"),
        ("7", "口径检验", "更换保守、核心、扩展AI口径，看结论是否稳定", "稳健性检验"),
    ]
    table = doc.add_table(rows=1, cols=4)
    set_table_widths(table, [1.0, 3.0, 9.0, 3.0])
    headers = ["步骤", "方法", "回答的问题", "定位"]
    for i, h in enumerate(headers):
        set_cell(table.rows[0].cells[i], h, fill=BLUE, bold=True, color="FFFFFF", align="center")
    for idx, name, question, role in rows:
        cells = table.add_row().cells
        set_cell(cells[0], idx, align="center", bold=True, color=BLUE)
        set_cell(cells[1], name, align="center", bold=True)
        set_cell(cells[2], question)
        set_cell(cells[3], role, align="center", bold=True, color=GREEN if "机制" in role else BLUE)


def add_formula_box(doc):
    add_callout(
        doc,
        "核心调节效应模型",
        "I_it = α + β1E_it + β2F_it + β3(E_it × F_it) + Controls_it + μ_i + λ_t + ε_it\n"
        "若 β3 > 0 且显著，说明金融支持越强，AI教育供给对AI产业发展的促进关系越强。\n"
        "注意：若没有更强识别设计，应表述为“正向调节作用”，而不是强因果意义上的“证明提升效率”。",
        fill=LIGHT_GREEN,
        border=GREEN,
    )


def build_doc():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.6)
    section.left_margin = Cm(1.9)
    section.right_margin = Cm(1.9)

    styles = doc.styles
    styles["Normal"].font.name = "微软雅黑"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    styles["Normal"].font.size = Pt(11)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("CCD、障碍度与机制检验的逻辑关系说明")
    set_cn_font(run, size=20, bold=True, color=BLUE)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("基于论文方法链条的讨论备忘 | 2026年5月15日")
    set_cn_font(run, size=10.5, color=GRAY)

    add_callout(
        doc,
        "核心结论",
        "这些方法可以放在同一篇论文里，但不能都写成“检验”。\n"
        "最稳的定位是：CCD = 描述性测度；障碍度 = 短板诊断；固定效应 = 经验关联检验；调节效应 = 机制检验；口径检验 = 稳健性检验。",
        fill=LIGHT_GOLD,
        border=GOLD,
    )

    add_heading(doc, "一、完整逻辑链", 1)
    add_flow_table(doc)

    add_heading(doc, "二、CCD与障碍度是否存在逻辑断裂？", 1)
    add_para(doc, "严格说，CCD和障碍度不是同一个问题的因果链条。CCD回答“协调水平如何”，障碍度回答“在既定指标体系下，哪些指标或子系统相对拖累更明显”。")
    add_para(doc, "因此，障碍度不能被表述为“证明存在短板”的方法。因为无论CCD高低，障碍度都能算出排序。它真正的价值是排序、诊断和政策优先级判断。")
    add_callout(
        doc,
        "推荐写法",
        "CCD用于识别金融、教育、产业三系统的协调水平及其变化；障碍度模型进一步在指标层面识别制约系统协调提升的主要因素，为解释CCD结果和提出政策建议提供依据。\n"
        "需要注意：障碍度并不是数学上分解CCD公式本身，而是解释哪些指标拖累了系统发展，从而可能影响协调水平。",
        fill=LIGHT_BLUE,
        border=BLUE,
    )

    add_heading(doc, "三、如果I系统是短板，后续固定效应和调节效应是否还有意义？", 1)
    add_para(doc, "有意义，而且论文主线会更清楚。若障碍度显示AI产业发展 I 是短板，后续模型正好可以解释：AI教育供给 E 和金融支持 F 是否有助于改善 I，以及 F 是否强化 E -> I 的转化。")
    add_para(doc, "此时逻辑不是“障碍度证明金融支持影响I”，而是“障碍度提出问题，回归模型解释问题”。")
    add_formula_box(doc)

    add_heading(doc, "四、什么时候后续模型才会真的失去意义？", 1)
    add_para(doc, "真正危险的情况不是 I 是短板，而是 I 几乎没有时间变化或地区差异。若AI产业发展指数在样本中长期很低且几乎不变，固定效应模型就很难估计出有意义的系数。")
    add_para(doc, "因此，后续必须先确认样本单位。如果 i 是省份，应采用省际面板；如果 i 是福建省内城市，需要市级AI教育、金融和产业数据；如果只有福建单省时间序列，则不能直接套用双向固定效应模型。")

    add_heading(doc, "五、可直接写入论文的方法衔接表述", 1)
    add_callout(
        doc,
        "论文表述建议",
        "本文首先基于保守、核心和扩展三类AI识别口径，构建金融支持、AI教育供给和AI产业发展三个子系统评价指标，并测算各子系统发展水平。在此基础上，采用耦合协调度模型刻画三系统协调状态，进一步运用障碍度模型识别制约协调水平提升的主要指标和短板系统。\n"
        "若障碍度结果显示AI产业发展系统是主要短板，则本文进一步以AI产业发展水平为被解释变量，检验AI教育供给和金融支持是否能够解释产业系统的变化，并重点考察金融支持是否在AI教育供给向AI产业发展转化过程中发挥正向调节作用。\n"
        "最后，通过替换AI识别口径、滞后变量和权重方法进行稳健性检验，以考察主要结论是否依赖特定指标设定。",
        fill=LIGHT_GREEN,
        border=GREEN,
    )

    add_heading(doc, "六、一句话概括", 1)
    add_callout(
        doc,
        "最简逻辑",
        "CCD回答“协调水平如何”；障碍度回答“短板在哪里”；固定效应回答“E和F是否与I相关”；调节效应回答“F是否强化E向I的转化”；口径检验回答“结论是否依赖AI定义”。",
        fill=LIGHT_RED,
        border=RED,
    )

    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("论文方法逻辑讨论备忘 | 仅供修改与导师沟通使用")
    set_cn_font(run, size=9, color=GRAY)

    doc.save(OUT_DOCX)
    return OUT_DOCX


if __name__ == "__main__":
    print(build_doc())

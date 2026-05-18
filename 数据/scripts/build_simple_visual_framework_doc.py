from pathlib import Path
import math

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


THESIS_DIR = Path(r"C:\金融专硕\论文\硕士毕业论文_Obsidian_MD统一整理_20260512")
OUT_DIR = THESIS_DIR / "06_导师沟通材料"
ASSET_DIR = OUT_DIR / "visual_assets_simple_framework"
OUT_DOCX = OUT_DIR / "简易论文框架_图示版_20260514.docx"

FONT_NORMAL = r"C:\Windows\Fonts\msyh.ttc"
FONT_BOLD = r"C:\Windows\Fonts\simhei.ttf"

COLORS = {
    "navy": "#1f4e79",
    "blue": "#2f75b5",
    "sky": "#eaf3fb",
    "pale": "#f6fbff",
    "gold": "#d9a441",
    "gold_pale": "#fff7df",
    "green": "#2e7d60",
    "green_pale": "#eaf7f0",
    "red": "#b55454",
    "red_pale": "#fff0ef",
    "gray": "#6b7280",
    "line": "#b9c7d6",
    "dark": "#1f2933",
    "white": "#ffffff",
}


def font(size, bold=False):
    return ImageFont.truetype(FONT_BOLD if bold else FONT_NORMAL, size)


def hex_color(value):
    value = value.lstrip("#")
    return tuple(int(value[i:i + 2], 16) for i in (0, 2, 4))


def wrap_text(draw, text, fnt, max_width):
    lines = []
    for para in text.split("\n"):
        current = ""
        for char in para:
            test = current + char
            width = draw.textbbox((0, 0), test, font=fnt)[2]
            if width <= max_width or not current:
                current = test
            else:
                lines.append(current)
                current = char
        lines.append(current)
    return lines


def draw_text_box(draw, xy, text, fill, outline, title=None, title_fill=None,
                  text_color=None, title_color=None, radius=24,
                  padding=22, body_size=34, title_size=36, align="center"):
    x1, y1, x2, y2 = xy
    draw.rounded_rectangle(xy, radius=radius, fill=hex_color(fill), outline=hex_color(outline), width=3)
    inner_y = y1 + padding
    if title:
        h = 58
        draw.rounded_rectangle((x1, y1, x2, y1 + h), radius=radius, fill=hex_color(title_fill or outline))
        draw.rectangle((x1, y1 + h - radius, x2, y1 + h), fill=hex_color(title_fill or outline))
        tf = font(title_size, bold=True)
        tw = draw.textbbox((0, 0), title, font=tf)[2]
        draw.text((x1 + (x2 - x1 - tw) / 2, y1 + 10), title, fill=hex_color(title_color or COLORS["white"]), font=tf)
        inner_y = y1 + h + padding - 4
    bf = font(body_size)
    lines = wrap_text(draw, text, bf, x2 - x1 - padding * 2)
    line_h = body_size + 12
    total_h = len(lines) * line_h
    start_y = inner_y
    if align == "center":
        start_y = max(inner_y, y1 + (y2 - y1 - total_h) / 2 + (30 if title else 0))
    for idx, line in enumerate(lines):
        bbox = draw.textbbox((0, 0), line, font=bf)
        if align == "left":
            tx = x1 + padding
        else:
            tx = x1 + (x2 - x1 - (bbox[2] - bbox[0])) / 2
        draw.text((tx, start_y + idx * line_h), line, fill=hex_color(text_color or COLORS["dark"]), font=bf)


def arrow(draw, start, end, color=None, width=7):
    color = hex_color(color or COLORS["blue"])
    draw.line([start, end], fill=color, width=width)
    x1, y1 = start
    x2, y2 = end
    angle = math.atan2(y2 - y1, x2 - x1)
    length = 24
    spread = 0.55
    points = [
        (x2, y2),
        (x2 - length * math.cos(angle - spread), y2 - length * math.sin(angle - spread)),
        (x2 - length * math.cos(angle + spread), y2 - length * math.sin(angle + spread)),
    ]
    draw.polygon(points, fill=color)


def base_canvas(title, subtitle=None):
    img = Image.new("RGB", (1800, 980), hex_color(COLORS["white"]))
    draw = ImageDraw.Draw(img)
    draw.rectangle((0, 0, 1800, 118), fill=hex_color(COLORS["navy"]))
    draw.text((64, 30), title, fill=hex_color(COLORS["white"]), font=font(46, bold=True))
    if subtitle:
        draw.text((64, 84), subtitle, fill=hex_color("#dbeafe"), font=font(22))
    return img, draw


def save_overview():
    img, draw = base_canvas("论文整体框架一页图", "从区域问题到金融调节机制")
    nodes = [
        ("现实背景", "福建AI产业集聚、资本支持与成果转化仍有提升空间", COLORS["red_pale"], COLORS["red"]),
        ("理论起点", "教育-产业耦合\n区域创新系统\n金融发展与创新", COLORS["sky"], COLORS["blue"]),
        ("研究对象", "AI教育供给 E\n金融支持 F\nAI产业发展 I", COLORS["green_pale"], COLORS["green"]),
        ("核心模型", "金融支持是否强化\nE -> I 的转化关系\n关注 β3", COLORS["gold_pale"], COLORS["gold"]),
        ("政策落点", "科技金融\n产教融合\n成果转化\n应用场景", COLORS["pale"], COLORS["navy"]),
    ]
    xs = [70, 410, 750, 1090, 1430]
    y = 190
    for i, (title, text, fill, outline) in enumerate(nodes):
        draw_text_box(draw, (xs[i], y, xs[i] + 300, y + 250), text, fill, outline, title=title,
                      title_fill=outline, body_size=30, title_size=32)
        if i < len(nodes) - 1:
            arrow(draw, (xs[i] + 306, y + 125), (xs[i + 1] - 18, y + 125), COLORS["blue"])
    draw_text_box(
        draw,
        (90, 535, 1710, 790),
        "主线收窄：不要把论文写成单纯测算 D 值的“三系统耦合协调研究”，而应围绕“金融支持是否强化 AI 教育供给向 AI 产业发展的转化效率”展开。\n\n角色分工：CCD/障碍度负责描述现状与短板；调节效应模型支撑金融主线；第三产业AI应用强度作为产业吸纳机制或拓展变量。",
        COLORS["gold_pale"],
        COLORS["gold"],
        title="最重要的研究定位",
        title_fill=COLORS["gold"],
        align="left",
        body_size=30,
    )
    draw_text_box(
        draw,
        (90, 830, 1710, 925),
        "明天优先问导师：是否同意把题目从“三系统耦合协调测度”收窄为“金融支持调节AI教育供给向AI产业转化”？",
        COLORS["sky"],
        COLORS["blue"],
        body_size=31,
    )
    path = ASSET_DIR / "01_overview.png"
    img.save(path)
    return path


def save_mechanism():
    img, draw = base_canvas("核心机制图", "AI教育供给、金融支持与AI产业发展")
    draw_text_box(draw, (90, 290, 470, 570), "AI相关专业\nAI课程与培养方案\n硕博点与师资\n高校论文与专利", COLORS["sky"], COLORS["blue"], title="AI教育供给 E", title_fill=COLORS["blue"], body_size=32)
    draw_text_box(draw, (1330, 290, 1710, 570), "AI企业数量\nAI岗位需求\n企业专利与产品\n技术合同与场景", COLORS["green_pale"], COLORS["green"], title="AI产业发展 I", title_fill=COLORS["green"], body_size=32)
    arrow(draw, (500, 430), (1300, 430), COLORS["navy"], width=10)
    draw.text((760, 372), "教育成果产业化路径", fill=hex_color(COLORS["navy"]), font=font(38, bold=True))
    draw.text((820, 430), "人才 / 知识 / 专利 / 项目", fill=hex_color(COLORS["gray"]), font=font(30))
    draw_text_box(draw, (690, 135, 1110, 305), "资本供给 / 风险分担\n创新筛选 / 治理约束", COLORS["gold_pale"], COLORS["gold"], title="金融支持 F", title_fill=COLORS["gold"], body_size=25)
    arrow(draw, (900, 305), (900, 405), COLORS["gold"], width=8)
    draw.text((1015, 332), "调节作用：β3 > 0 ?", fill=hex_color(COLORS["gold"]), font=font(30, bold=True))
    draw_text_box(draw, (570, 615, 1230, 770), "第三产业AI应用强度\n信息服务、金融、医疗、教育、政务等数据密集场景\n作为产业吸纳机制或异质性变量", COLORS["pale"], COLORS["navy"], body_size=30)
    arrow(draw, (900, 615), (900, 470), COLORS["navy"], width=6)
    draw_text_box(draw, (100, 815, 1700, 920), "解释逻辑：AI教育供给提供人力资本和知识产出；金融支持缓解技术商业化融资约束；产业部门通过场景和岗位需求吸纳AI成果。", COLORS["green_pale"], COLORS["green"], body_size=31)
    path = ASSET_DIR / "02_mechanism.png"
    img.save(path)
    return path


def save_indicators():
    img, draw = base_canvas("指标体系图", "不追求机械一一对应，而是做功能链条对应")
    headers = [
        ("教育系统 E", COLORS["blue"], COLORS["sky"], ["规模基础：AI专业、招生规模", "高层次能力：硕博点、师资", "知识产出：论文、专利", "转化连接：产业学院、实验室", "应用吸纳：AI+X课程、毕业去向"]),
        ("产业系统 I", COLORS["green"], COLORS["green_pale"], ["规模基础：AI企业、岗位", "高层次能力：高企、专精特新", "技术产出：企业专利、产品", "转化连接：技术合同、示范场景", "应用吸纳：第三产业AI强度"]),
        ("金融系统 F", COLORS["gold"], COLORS["gold_pale"], ["一般金融：金融业/GDP、存贷款", "科技金融：科技贷款、数字金融", "风险资本：VC/PE、融资事件", "政府引导：基金、补贴、担保", "转化支持：知识产权质押、概念验证"]),
    ]
    x0s = [85, 635, 1185]
    for idx, (head, dark, light, items) in enumerate(headers):
        x1 = x0s[idx]
        draw.rounded_rectangle((x1, 170, x1 + 530, 835), radius=28, fill=hex_color(light), outline=hex_color(dark), width=4)
        draw.rounded_rectangle((x1, 170, x1 + 530, 240), radius=28, fill=hex_color(dark))
        draw.rectangle((x1, 210, x1 + 530, 240), fill=hex_color(dark))
        draw.text((x1 + 155, 188), head, fill=hex_color(COLORS["white"]), font=font(36, bold=True))
        for j, item in enumerate(items):
            y = 280 + j * 100
            draw.rounded_rectangle((x1 + 35, y, x1 + 495, y + 70), radius=18, fill=hex_color(COLORS["white"]), outline=hex_color("#d6dee8"), width=2)
            draw.text((x1 + 60, y + 17), item, fill=hex_color(COLORS["dark"]), font=font(28))
    draw_text_box(draw, (120, 870, 1680, 935), "导师重点确认：公开可得指标、替代变量，以及金融变量是“一般金融环境”还是“AI专项金融支持”。", COLORS["red_pale"], COLORS["red"], body_size=28)
    path = ASSET_DIR / "03_indicators.png"
    img.save(path)
    return path


def save_methods():
    img, draw = base_canvas("研究方法路线图", "描述测度服务现状判断，调节效应服务金融主线")
    steps = [
        ("数据采集", "年鉴 / 教育部\n专利 / 企业 / 融资"),
        ("指标处理", "AI口径\n标准化 / 赋权"),
        ("子系统指数", "E 教育\nF 金融\nI 产业"),
        ("描述测度", "CCD\n障碍度"),
        ("机制检验", "固定效应\nE × F 调节"),
        ("稳健性", "口径替换\n滞后项 / 等权重"),
        ("政策建议", "科技金融\n产教融合"),
    ]
    x = 80
    y = 235
    for i, (title, body) in enumerate(steps):
        fill = COLORS["sky"] if i < 3 else COLORS["gold_pale"] if i < 5 else COLORS["green_pale"]
        outline = COLORS["blue"] if i < 3 else COLORS["gold"] if i < 5 else COLORS["green"]
        draw_text_box(draw, (x, y, x + 210, y + 210), body, fill, outline, title=title, title_fill=outline, body_size=25, title_size=28)
        if i < len(steps) - 1:
            arrow(draw, (x + 216, y + 105), (x + 270, y + 105), COLORS["navy"], width=6)
        x += 245
    draw_text_box(draw, (130, 535, 850, 745), "描述性部分：三系统发展水平与短板。\n输出：子系统指数、CCD结果、障碍指标。", COLORS["pale"], COLORS["navy"], title="CCD与障碍度", title_fill=COLORS["navy"], align="left", body_size=28)
    draw_text_box(draw, (950, 535, 1670, 745), "解释性部分：金融是否强化 E -> I。\n核心：I_it = α + β1E_it + β2F_it + β3(E×F) + Controls + FE。", COLORS["green_pale"], COLORS["green"], title="调节效应模型", title_fill=COLORS["green"], align="left", body_size=27)
    draw_text_box(draw, (130, 815, 1670, 925), "风险提示：若使用固定效应模型，必须明确 i 是省份还是福建省内城市；若只做福建单省时间序列，不能直接套用双向固定效应。", COLORS["red_pale"], COLORS["red"], body_size=31)
    path = ASSET_DIR / "04_methods.png"
    img.save(path)
    return path


def save_chapters():
    img, draw = base_canvas("章节安排图", "建议七章结构：先定边界，再测度，最后解释机制")
    chapters = [
        ("第一章", "导论", "问题提出\n研究意义\n技术路线"),
        ("第二章", "文献综述", "教育-产业耦合\n金融与创新\nAI技能供给"),
        ("第三章", "理论机制", "E -> I\nF调节\n产业吸纳"),
        ("第四章", "指标与方法", "AI口径\n指标体系\n模型设定"),
        ("第五章", "现状测度", "子系统指数\nCCD\n障碍度"),
        ("第六章", "机制检验", "固定效应\n调节效应\n稳健性"),
        ("第七章", "结论建议", "科技金融\n产教融合\n成果转化"),
    ]
    y_positions = [175, 285, 395, 505, 615, 725, 835]
    for i, (num, title, body) in enumerate(chapters):
        y = y_positions[i]
        color = COLORS["blue"] if i < 2 else COLORS["gold"] if i < 4 else COLORS["green"] if i < 6 else COLORS["navy"]
        draw.rounded_rectangle((95, y, 410, y + 80), radius=18, fill=hex_color(color))
        draw.text((128, y + 20), f"{num}  {title}", fill=hex_color(COLORS["white"]), font=font(28, bold=True))
        draw.rounded_rectangle((440, y, 1045, y + 80), radius=18, fill=hex_color(COLORS["white"]), outline=hex_color(color), width=3)
        draw.text((470, y + 19), body.replace("\n", "  /  "), fill=hex_color(COLORS["dark"]), font=font(28))
        if i < len(chapters) - 1:
            arrow(draw, (252, y + 82), (252, y + 108), color, width=5)
    draw_text_box(
        draw,
        (1110, 175, 1690, 790),
        "1. 题目是否正式收窄为金融调节机制？\n\n2. 样本采用省际面板、福建省内城市面板，还是福建单省描述？\n\n3. AI指标采用保守、核心、扩展三类口径是否认可？\n\n4. CCD是描述性测度，调节效应是主模型，是否认可？\n\n5. 受限数据库能否用公开替代变量先行？",
        COLORS["gold_pale"],
        COLORS["gold"],
        title="明天请导师确认",
        title_fill=COLORS["gold"],
        align="left",
        body_size=30,
    )
    draw_text_box(draw, (1110, 830, 1690, 925), "一句话：金融支持是否强化AI教育供给向AI产业转化。", COLORS["sky"], COLORS["blue"], body_size=27)
    path = ASSET_DIR / "05_chapters.png"
    img.save(path)
    return path


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill.replace("#", ""))
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, color="1f2933", size=12):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "微软雅黑"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_title(paragraph, text, size=22, color="1f4e79"):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(text)
    run.bold = True
    run.font.name = "微软雅黑"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)


def add_note_table(doc, text, fill="FFF7DF", border="D9A441"):
    table = doc.add_table(rows=1, cols=1)
    table.autofit = True
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.font.name = "微软雅黑"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(31, 41, 51)
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for side in ["top", "left", "bottom", "right"]:
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "8")
        el.set(qn("w:color"), border)
        borders.append(el)
    tc_pr.append(borders)


def setup_document():
    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(11.69)
    section.page_height = Inches(8.27)
    section.top_margin = Cm(1.05)
    section.bottom_margin = Cm(1.05)
    section.left_margin = Cm(1.05)
    section.right_margin = Cm(1.05)
    styles = doc.styles
    styles["Normal"].font.name = "微软雅黑"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    styles["Normal"].font.size = Pt(11)
    return doc


def add_image_page(doc, title, img_path, note=None, page_break=True):
    if page_break and len(doc.paragraphs) > 1:
        doc.add_page_break()
    p = doc.add_paragraph()
    add_title(p, title)
    p.paragraph_format.space_after = Pt(6)
    doc.add_picture(str(img_path), width=Inches(10.65))


def build_doc():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    images = [
        save_overview(),
        save_mechanism(),
        save_indicators(),
        save_methods(),
        save_chapters(),
    ]

    doc = setup_document()
    add_image_page(
        doc,
        "简易论文框架：图示版",
        images[0],
        "用途：快速确认论文方向、指标、样本和模型。",
        page_break=False,
    )
    add_image_page(
        doc,
        "一、核心机制：金融支持作为转化器",
        images[1],
        "汇报句：AI教育供给提供人才和知识，金融支持降低商业化约束，产业场景吸纳成果。",
    )
    add_image_page(
        doc,
        "二、指标体系：三系统与功能链条",
        images[2],
        "关键提醒：指标不宜机械一一对应，应先区分教育供给、知识产出、产业发展、金融支持和产业吸纳机制。",
    )
    add_image_page(
        doc,
        "三、研究方法：描述测度与机制检验分工",
        images[3],
        "导师确认点：若采用固定效应模型，必须先确定样本单位。",
    )
    add_image_page(
        doc,
        "四、章节安排与明日确认清单",
        images[4],
        "建议结论：CCD/障碍度描述短板，调节效应模型支撑金融主线。",
    )

    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("硕士论文简易框架图示版 | 仅供导师讨论")
    run.font.name = "微软雅黑"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(107, 114, 128)

    doc.save(OUT_DOCX)
    return OUT_DOCX


if __name__ == "__main__":
    print(build_doc())

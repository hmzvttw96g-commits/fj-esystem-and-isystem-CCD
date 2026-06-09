from __future__ import annotations

import csv
import json
import logging
import math
import re
import textwrap
import zipfile
from collections import Counter, defaultdict
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, Iterable, List, Sequence, Tuple

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from openpyxl import Workbook, load_workbook
from openpyxl.styles import Alignment, Font, PatternFill
from openpyxl.utils import get_column_letter
from PIL import Image, ImageDraw, ImageFont
from pypdf import PdfReader
from reportlab.lib.pagesizes import A4
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.cidfonts import UnicodeCIDFont
from reportlab.pdfgen import canvas


ROOT = Path(__file__).resolve().parents[1]
RAW_DIR = ROOT / "data" / "raw"
PROCESSED_DIR = ROOT / "data" / "processed"
REPORTS_DIR = ROOT / "reports"
CHARTS_DIR = REPORTS_DIR / "charts"
GENERATED_DIR = ROOT / "generated_papers"
LOGS_DIR = ROOT / "logs"
TEMP_DIR = ROOT / "temp"

for _path in [PROCESSED_DIR, REPORTS_DIR, CHARTS_DIR, GENERATED_DIR, LOGS_DIR, TEMP_DIR]:
    _path.mkdir(parents=True, exist_ok=True)


QUESTION_FIELDS = [
    "source_type",
    "school",
    "year",
    "file_name",
    "question_id",
    "question_type",
    "question_text",
    "options",
    "score",
    "answer",
    "solution",
    "knowledge_point",
    "detailed_knowledge_point",
    "chapter",
    "difficulty",
    "has_image",
    "has_formula",
]


KNOWLEDGE_KEYWORDS: Dict[str, List[str]] = {
    "集合与逻辑": ["集合", "子集", "真子集", "空集", "并集", "交集", "补集", "元素", "区间", "充分", "必要", "充要", "命题"],
    "不等式": ["不等式", "解集", "大于", "小于", "范围", "≤", "≥", ">=", "<=", "高于", "低于"],
    "函数": ["函数", "定义域", "值域", "f(", "f（", "单调", "奇函数", "偶函数", "分段", "图像"],
    "指数与对数": ["指数", "对数", "log", "lg", "ln", "幂", "根式"],
    "三角函数": ["sin", "cos", "tan", "正弦", "余弦", "正切", "弧度", "角", "周期"],
    "数列": ["数列", "等差", "等比", "第", "a4", "a5", "项"],
    "向量": ["向量", "相反向量", "内积", "数量积", "AB", "BC", "→", "平行向量"],
    "直线与圆": ["直线", "斜率", "圆", "圆心", "半径", "方程", "坐标", "x轴", "y轴"],
    "立体几何": ["长方体", "正方体", "圆柱", "圆锥", "球", "体积", "侧面", "平面", "异面", "空间"],
    "概率统计": ["概率", "抽样", "事件", "统计", "样本", "任选", "选法", "编号"],
    "应用与建模": ["实际", "模型", "生产", "生活", "职业", "情境", "数据处理"],
}

DETAILED_KNOWLEDGE_RULES: List[Tuple[str, str, List[str]]] = [
    ("集合交并补运算", "集合与逻辑", ["交集", "并集", "补集", "∩", "∪", "全集", "U A", "U Að", "A∩B"]),
    ("集合表示与元素关系", "集合与逻辑", ["集合", "元素", "属于", "∈", "空集", "有限集", "无限集", "自然数", "正整数"]),
    ("区间与描述法", "集合与逻辑", ["区间", "{x|", "描述法", "用区间表示", "解集用区间"]),
    ("充分必要条件", "集合与逻辑", ["充分", "必要", "充要"]),
    ("一元一次不等式", "不等式", ["不等式", "3x", "2x", "解集"]),
    ("一元二次不等式", "不等式", ["一元二次不等式", "(x", ")( ", "二次", "判别式"]),
    ("绝对值不等式", "不等式", ["|x", "绝对值"]),
    ("函数求值", "函数", ["f(", "f（", "求 f", "f(3)", "f(-", "f(0)"]),
    ("函数定义域", "函数", ["定义域", "函数的定义域"]),
    ("函数奇偶性", "函数", ["奇函数", "偶函数", "奇偶", "既不是奇函数"]),
    ("函数单调性", "函数", ["增函数", "减函数", "单调"]),
    ("分段函数", "函数", ["分段函数", "分段"]),
    ("一次函数模型", "函数", ["一次函数", "函数关系式", "y关于x", "收益", "费用", "投资时间", "温度"]),
    ("指数式与对数式互化", "指数与对数", ["指数式", "化为对数式", "log", "对数式"]),
    ("对数运算", "指数与对数", ["log", "lg", "ln", "对数"]),
    ("指数幂运算", "指数与对数", ["指数", "幂", "^", "2^", "a^"]),
    ("特殊角三角函数", "三角函数", ["sin30", "cos60", "sin90", "特殊角"]),
    ("三角函数符号与终边", "三角函数", ["终边", "角", "sin", "cos", "tan", "象限"]),
    ("等差数列", "数列", ["等差", "公差"]),
    ("等比数列", "数列", ["等比", "公比"]),
    ("数列通项与递推", "数列", ["通项", "递推", "前项和", "数列", "a_n", "an", "第4项", "第5项"]),
    ("向量坐标运算", "向量", ["向量", "坐标", "a=(", "b=(", "a+b", "数量积", "内积"]),
    ("向量几何表示", "向量", ["AB", "BC", "DB", "平行四边形", "相反向量", "位移"]),
    ("直线斜率", "直线与圆", ["斜率", "直线 y=", "直线的斜率"]),
    ("直线位置关系", "直线与圆", ["平行", "垂直", "位置关系", "l1", "l2"]),
    ("圆的标准方程", "直线与圆", ["圆", "圆心", "半径", "标准方程", "(x-", "(y"]),
    ("直线与圆综合", "直线与圆", ["直线", "圆", "相切", "弦", "交于", "切线"]),
    ("立体几何体积", "立体几何", ["体积", "长方体", "正方体", "圆柱", "圆锥", "球"]),
    ("立体几何视图与位置", "立体几何", ["俯视图", "主视图", "异面", "正方体", "棱柱", "棱锥", "平面"]),
    ("计数原理", "概率统计", ["选法", "任选", "搭配", "不同", "可以炒成", "共有"]),
    ("抽样与频率分布", "概率统计", ["抽样", "系统抽样", "分层抽样", "频率分布", "样本容量"]),
    ("概率基础", "概率统计", ["概率", "随机抽取", "事件"]),
    ("应用函数建模", "应用与建模", ["实际", "模型", "费用", "收益", "温度", "投资", "生产", "生活"]),
]

CHAPTER_MAP = {
    "集合与逻辑": "基础知识",
    "不等式": "基础知识",
    "函数": "函数",
    "指数与对数": "函数",
    "三角函数": "函数",
    "数列": "数列",
    "向量": "几何与向量",
    "直线与圆": "几何与向量",
    "立体几何": "几何与向量",
    "概率统计": "概率统计",
    "应用与建模": "应用与建模",
    "其他": "其他",
}


def setup_logger(name: str, log_file: Path) -> logging.Logger:
    log_file.parent.mkdir(parents=True, exist_ok=True)
    logger = logging.getLogger(name)
    logger.setLevel(logging.INFO)
    logger.handlers.clear()
    fmt = logging.Formatter("%(asctime)s | %(levelname)s | %(message)s")
    fh = logging.FileHandler(log_file, encoding="utf-8")
    fh.setFormatter(fmt)
    sh = logging.StreamHandler()
    sh.setFormatter(fmt)
    logger.addHandler(fh)
    logger.addHandler(sh)
    return logger


def now_str() -> str:
    return datetime.now().strftime("%Y-%m-%d %H:%M:%S")


def clean_text(text: str) -> str:
    if not text:
        return ""
    text = text.encode("utf-8", "replace").decode("utf-8", "replace")
    text = text.replace("\ufeff", "")
    watermark_patterns = [
        r"B\s*站[，,、]\s*小红书\s*ID[:：][^\n]+",
        r"B站[:：]?[^\n]*学数学的小子曾帅",
        r"知源学考",
    ]
    for pat in watermark_patterns:
        text = re.sub(pat, "", text)
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def read_docx_text(path: Path) -> Tuple[str, Dict[str, Any]]:
    doc = Document(str(path))
    parts: List[str] = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            vals = [cell.text.strip() for cell in row.cells]
            if any(vals):
                parts.append("\t".join(vals))
    meta = {"pages": None, "chars": sum(len(p) for p in parts), "ocr_failed": False}
    return clean_text("\n".join(parts)), meta


def read_pdf_text(path: Path) -> Tuple[str, Dict[str, Any]]:
    reader = PdfReader(str(path))
    parts: List[str] = []
    page_lengths: List[int] = []
    for page in reader.pages:
        page_text = page.extract_text() or ""
        page_text = clean_text(page_text)
        page_lengths.append(len(page_text))
        parts.append(page_text)
    text = clean_text("\n".join(parts))
    meta = {
        "pages": len(reader.pages),
        "chars": len(text),
        "page_lengths": page_lengths,
        "ocr_failed": len(text) < max(100, len(reader.pages) * 40),
    }
    return text, meta


def read_xlsx_text(path: Path) -> Tuple[str, Dict[str, Any]]:
    wb = load_workbook(path, read_only=True, data_only=True)
    parts: List[str] = []
    cells = 0
    for ws in wb.worksheets:
        parts.append(f"【工作表】{ws.title}")
        for row in ws.iter_rows(values_only=True):
            vals = [str(v) for v in row if v is not None and str(v).strip()]
            if vals:
                cells += len(vals)
                parts.append("\t".join(vals))
    text = clean_text("\n".join(parts))
    return text, {"chars": len(text), "cells": cells, "ocr_failed": False}


def read_html_text(path: Path) -> Tuple[str, Dict[str, Any]]:
    text = path.read_text(encoding="utf-8", errors="replace")
    text = re.sub(r"<script.*?</script>", " ", text, flags=re.S | re.I)
    text = re.sub(r"<style.*?</style>", " ", text, flags=re.S | re.I)
    text = re.sub(r"<[^>]+>", " ", text)
    text = clean_text(text)
    return text, {"chars": len(text), "ocr_failed": False}


def read_supported_text(path: Path) -> Tuple[str, Dict[str, Any]]:
    suffix = path.suffix.lower()
    if suffix == ".docx":
        return read_docx_text(path)
    if suffix == ".pdf":
        return read_pdf_text(path)
    if suffix in {".xlsx", ".xlsm"}:
        return read_xlsx_text(path)
    if suffix in {".html", ".htm"}:
        return read_html_text(path)
    if suffix in {".txt", ".md"}:
        text = path.read_text(encoding="utf-8", errors="replace")
        return clean_text(text), {"chars": len(text), "ocr_failed": False}
    if suffix in {".png", ".jpg", ".jpeg", ".bmp", ".tif", ".tiff", ".webp"}:
        try:
            img = Image.open(path)
            return "", {"chars": 0, "image_size": img.size, "ocr_failed": True, "notes": "image_requires_ocr"}
        except Exception as exc:
            return "", {"chars": 0, "ocr_failed": True, "notes": str(exc)}
    return "", {"unsupported": True, "ocr_failed": True, "chars": 0}


def classify_knowledge(text: str) -> str:
    scores: Dict[str, int] = {}
    for kp, words in KNOWLEDGE_KEYWORDS.items():
        score = 0
        for w in words:
            score += text.lower().count(w.lower())
        if score:
            scores[kp] = score
    if not scores:
        return "其他"
    return max(scores.items(), key=lambda x: x[1])[0]


def classify_detailed_knowledge(text: str, broad: str | None = None) -> str:
    broad = broad or classify_knowledge(text)
    lowered = text.lower()
    best: Tuple[int, str] = (0, "")
    for detail, detail_broad, words in DETAILED_KNOWLEDGE_RULES:
        score = 0
        for word in words:
            if word.lower() in lowered:
                score += 1
        if broad and detail_broad == broad:
            score += 1
        if score > best[0]:
            best = (score, detail)
    if best[0] > 0:
        return best[1]
    fallback = {
        "集合与逻辑": "集合与逻辑基础",
        "不等式": "不等式基础",
        "函数": "函数基础",
        "指数与对数": "指数对数基础",
        "三角函数": "三角函数基础",
        "数列": "数列基础",
        "向量": "向量基础",
        "直线与圆": "解析几何基础",
        "立体几何": "立体几何基础",
        "概率统计": "概率统计基础",
        "应用与建模": "应用建模基础",
    }
    return fallback.get(str(broad), "其他待人工复核")


def infer_chapter(knowledge_point: str) -> str:
    return CHAPTER_MAP.get(knowledge_point, "其他")


def infer_question_type(heading: str, qid: str, text: str) -> str:
    if "选择" in heading:
        return "选择题"
    if "填空" in heading:
        return "填空题"
    if "解答" in heading or "计算" in heading:
        return "解答题"
    if re.search(r"\b[A-D][.．、]", text):
        return "选择题"
    if "____" in text or "________" in text:
        return "填空题"
    return "解答题" if len(text) > 120 else "其他"


def infer_score(heading: str, qtype: str, text: str) -> int:
    header = heading + " " + text[:120]
    m = re.search(r"每小题\s*(\d+)\s*分", header)
    if m:
        return int(m.group(1))
    m = re.search(r"共\s*(\d+)\s*分", header)
    if qtype == "解答题" and m:
        return int(m.group(1))
    if qtype == "选择题":
        return 5
    if qtype == "填空题":
        return 5
    if qtype == "解答题":
        return 10
    return 0


def infer_difficulty(qtype: str, text: str, knowledge_point: str) -> str:
    hard_markers = ["证明", "综合", "分类讨论", "取值范围", "分段函数", "不等式模型", "应用"]
    medium_kp = {"三角函数", "向量", "直线与圆", "立体几何", "指数与对数", "函数"}
    if qtype == "解答题" or any(m in text for m in hard_markers):
        return "较难"
    if qtype == "填空题" or knowledge_point in medium_kp:
        return "中等"
    return "基础"


def extract_answer(text: str) -> str:
    for pat in [
        r"【参考答案】\s*([A-D]|[^。\n]{1,30})",
        r"答案[:：]\s*([A-D]|[^。\n]{1,30})",
        r"[（(]\s*([A-D])\s*[）)]",
    ]:
        m = re.search(pat, text)
        if m:
            return clean_text(m.group(1))
    return ""


def extract_options(text: str) -> str:
    opts = []
    for label in ["A", "B", "C", "D"]:
        m = re.search(rf"{label}[.．、]\s*(.*?)(?=\s+[A-D][.．、]|\n|$)", text, flags=re.S)
        if m:
            opts.append(f"{label}. {clean_text(m.group(1))}")
    return " | ".join(opts)


def has_formula(text: str) -> bool:
    return bool(re.search(r"[=+\-×÷*/^≤≥<>√∠]|log|sin|cos|tan|f\(|\d+[xyabn]", text, flags=re.I))


def heading_positions(text: str) -> List[Tuple[int, str]]:
    headings = []
    for m in re.finditer(r"(一、[^。\n]{0,40}|二、[^。\n]{0,40}|三、[^。\n]{0,40}|选择题|填空题|解答题|等级卷|合格卷)", text):
        headings.append((m.start(), m.group(0)))
    return headings


def last_heading_before(headings: Sequence[Tuple[int, str]], pos: int) -> str:
    current = ""
    for hpos, htext in headings:
        if hpos <= pos:
            current = htext
        else:
            break
    return current


def parse_questions_from_text(
    text: str,
    *,
    source_type: str,
    school: str = "",
    year: str = "",
    file_name: str = "",
) -> List[Dict[str, Any]]:
    text = clean_text(text)
    if not text:
        return []
    starts = list(re.finditer(r"(?m)(?:^|\n)\s*(\d{1,2})[、.．]\s*", text))
    headings = heading_positions(text)
    rows: List[Dict[str, Any]] = []
    seen = Counter()
    for i, m in enumerate(starts):
        qid_raw = m.group(1)
        start = m.start()
        end = starts[i + 1].start() if i + 1 < len(starts) else min(len(text), start + 1500)
        q_text = clean_text(text[start:end])
        if len(q_text) < 8:
            continue
        heading = last_heading_before(headings, start)
        qtype = infer_question_type(heading, qid_raw, q_text)
        seen_key = f"{qid_raw}-{qtype}"
        seen[seen_key] += 1
        qid = qid_raw if seen[seen_key] == 1 else f"{qid_raw}-{seen[seen_key]}"
        kp = classify_knowledge(q_text)
        detailed_kp = classify_detailed_knowledge(q_text, kp)
        rows.append(
            {
                "source_type": source_type,
                "school": school,
                "year": str(year),
                "file_name": file_name,
                "question_id": qid,
                "question_type": qtype,
                "question_text": q_text[:1200],
                "options": extract_options(q_text),
                "score": infer_score(heading, qtype, q_text),
                "answer": extract_answer(q_text),
                "solution": clean_text(q_text.split("解：", 1)[1])[:800] if "解：" in q_text else "",
                "knowledge_point": kp,
                "detailed_knowledge_point": detailed_kp,
                "chapter": infer_chapter(kp),
                "difficulty": infer_difficulty(qtype, q_text, kp),
                "has_image": "如图" in q_text or "图" in q_text,
                "has_formula": has_formula(q_text),
            }
        )
    return rows


def syllabus_points_from_text(text: str, *, year: str, file_name: str) -> List[Dict[str, Any]]:
    text = clean_text(text)
    rows = []
    for kp in KNOWLEDGE_KEYWORDS:
        hits = sum(text.lower().count(w.lower()) for w in KNOWLEDGE_KEYWORDS[kp])
        if hits:
            rows.append(
                {
                    "source_type": "syllabus",
                    "school": "",
                    "year": year,
                    "file_name": file_name,
                    "question_id": f"SYL-{kp}",
                    "question_type": "考纲条目",
                    "question_text": kp,
                    "options": "",
                    "score": 0,
                    "answer": "",
                    "solution": "",
                    "knowledge_point": kp,
                    "detailed_knowledge_point": kp,
                    "chapter": infer_chapter(kp),
                    "difficulty": "",
                    "has_image": False,
                    "has_formula": False,
                }
            )
    return rows


def write_csv(rows: Sequence[Dict[str, Any]], path: Path, fields: Sequence[str] | None = None) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    if not rows and not fields:
        fields = []
    fields = list(fields or rows[0].keys())
    with path.open("w", newline="", encoding="utf-8-sig") as f:
        writer = csv.DictWriter(f, fieldnames=fields)
        writer.writeheader()
        for row in rows:
            writer.writerow({k: row.get(k, "") for k in fields})


def write_xlsx(sheets: Dict[str, Sequence[Dict[str, Any]]], path: Path) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    wb = Workbook()
    first = True
    for sheet_name, rows in sheets.items():
        ws = wb.active if first else wb.create_sheet()
        ws.title = sheet_name[:31]
        first = False
        fields: List[str]
        if rows:
            fields = list(rows[0].keys())
        else:
            fields = ["note"]
        ws.append(fields)
        for cell in ws[1]:
            cell.font = Font(bold=True, color="FFFFFF")
            cell.fill = PatternFill("solid", fgColor="366092")
            cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
        if rows:
            for row in rows:
                ws.append([row.get(f, "") for f in fields])
        for col_idx, field in enumerate(fields, 1):
            max_len = min(max([len(str(field))] + [len(str(r.get(field, ""))) for r in rows[:200]]) + 2, 48)
            ws.column_dimensions[get_column_letter(col_idx)].width = max(10, max_len)
        for row in ws.iter_rows():
            for cell in row:
                cell.alignment = Alignment(vertical="top", wrap_text=True)
        ws.freeze_panes = "A2"
        ws.auto_filter.ref = ws.dimensions
    wb.save(path)


def read_xlsx_rows(path: Path, sheet_name: str | None = None) -> List[Dict[str, Any]]:
    wb = load_workbook(path, data_only=True)
    ws = wb[sheet_name] if sheet_name else wb[wb.sheetnames[0]]
    rows = list(ws.iter_rows(values_only=True))
    if not rows:
        return []
    fields = [str(v) if v is not None else "" for v in rows[0]]
    out = []
    for vals in rows[1:]:
        out.append({fields[i]: vals[i] if i < len(vals) else "" for i in range(len(fields))})
    return out


def distribution(rows: Iterable[Dict[str, Any]], field: str, filter_fn=None) -> Counter:
    c = Counter()
    for r in rows:
        if filter_fn and not filter_fn(r):
            continue
        val = r.get(field) or "未识别"
        c[str(val)] += 1
    return c


def normalize_counter(counter: Counter) -> Dict[str, float]:
    total = sum(counter.values()) or 1
    return {k: v / total for k, v in counter.items()}


def overlap_rate(a: Iterable[str], b: Iterable[str]) -> float:
    sa = {x for x in a if x}
    sb = {x for x in b if x}
    if not sa or not sb:
        return 0.0
    return len(sa & sb) / len(sa | sb)


def l1_similarity(counter_a: Counter, counter_b: Counter) -> float:
    pa = normalize_counter(counter_a)
    pb = normalize_counter(counter_b)
    keys = set(pa) | set(pb)
    distance = sum(abs(pa.get(k, 0) - pb.get(k, 0)) for k in keys) / 2
    return max(0.0, 1 - distance)


def get_font(size: int = 22, bold: bool = False) -> ImageFont.FreeTypeFont | ImageFont.ImageFont:
    candidates = [
        "/System/Library/Fonts/STHeiti Medium.ttc",
        "/System/Library/Fonts/Supplemental/Songti.ttc",
        "/Library/Fonts/Arial Unicode.ttf",
    ]
    for path in candidates:
        if Path(path).exists():
            try:
                return ImageFont.truetype(path, size=size)
            except Exception:
                continue
    return ImageFont.load_default()


def draw_text(draw: ImageDraw.ImageDraw, xy, text: str, font, fill=(30, 30, 30), anchor=None):
    try:
        draw.text(xy, text, font=font, fill=fill, anchor=anchor)
    except UnicodeEncodeError:
        draw.text(xy, text.encode("utf-8", "replace").decode("utf-8"), font=font, fill=fill, anchor=anchor)


def save_bar_chart(labels: List[str], values: List[float], path: Path, title: str, unit: str = "") -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    w, h = 1200, 720
    img = Image.new("RGB", (w, h), "white")
    d = ImageDraw.Draw(img)
    title_font = get_font(34)
    label_font = get_font(18)
    small_font = get_font(16)
    draw_text(d, (w // 2, 34), title, title_font, anchor="ma")
    left, right, top, bottom = 110, 60, 110, 140
    max_v = max(values) if values else 1
    max_v = max_v or 1
    plot_w = w - left - right
    plot_h = h - top - bottom
    bar_gap = 14
    bar_w = max(22, (plot_w - bar_gap * (len(labels) + 1)) // max(1, len(labels)))
    colors = ["#3B82F6", "#10B981", "#F59E0B", "#EF4444", "#8B5CF6", "#14B8A6", "#64748B"]
    d.line((left, top, left, top + plot_h), fill="#333333", width=2)
    d.line((left, top + plot_h, left + plot_w, top + plot_h), fill="#333333", width=2)
    for i, (lab, val) in enumerate(zip(labels, values)):
        x0 = left + bar_gap + i * (bar_w + bar_gap)
        x1 = x0 + bar_w
        y1 = top + plot_h
        y0 = y1 - int((val / max_v) * (plot_h - 30))
        d.rectangle((x0, y0, x1, y1), fill=colors[i % len(colors)])
        draw_text(d, ((x0 + x1) // 2, y0 - 22), f"{val:.0f}{unit}", small_font, anchor="ma")
        wrapped = "\n".join(textwrap.wrap(str(lab), 6))
        draw_text(d, ((x0 + x1) // 2, y1 + 12), wrapped, label_font, anchor="ma")
    img.save(path)


def save_pie_chart(labels: List[str], values: List[float], path: Path, title: str) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    w, h = 1100, 720
    img = Image.new("RGB", (w, h), "white")
    d = ImageDraw.Draw(img)
    title_font = get_font(34)
    label_font = get_font(20)
    draw_text(d, (w // 2, 34), title, title_font, anchor="ma")
    total = sum(values) or 1
    colors = ["#2563EB", "#059669", "#D97706", "#DC2626", "#7C3AED", "#0891B2", "#4B5563", "#BE123C"]
    box = (120, 110, 620, 610)
    start = 0.0
    for i, val in enumerate(values):
        end = start + 360 * val / total
        d.pieslice(box, start, end, fill=colors[i % len(colors)], outline="white")
        start = end
    lx, ly = 690, 130
    for i, (lab, val) in enumerate(zip(labels, values)):
        y = ly + i * 52
        d.rectangle((lx, y, lx + 28, y + 28), fill=colors[i % len(colors)])
        draw_text(d, (lx + 42, y - 2), f"{lab}: {val / total:.1%}", label_font)
    img.save(path)


def save_heatmap(row_labels: List[str], col_labels: List[str], matrix: List[List[float]], path: Path, title: str) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    cell_w, cell_h = 160, 70
    left, top = 230, 120
    w = left + cell_w * len(col_labels) + 80
    h = top + cell_h * len(row_labels) + 110
    img = Image.new("RGB", (w, h), "white")
    d = ImageDraw.Draw(img)
    title_font = get_font(32)
    label_font = get_font(19)
    value_font = get_font(22)
    draw_text(d, (w // 2, 36), title, title_font, anchor="ma")
    max_v = max([v for row in matrix for v in row] or [1]) or 1
    for j, lab in enumerate(col_labels):
        draw_text(d, (left + j * cell_w + cell_w // 2, top - 38), lab, label_font, anchor="ma")
    for i, rlab in enumerate(row_labels):
        draw_text(d, (left - 14, top + i * cell_h + cell_h // 2), rlab, label_font, anchor="rm")
        for j, val in enumerate(matrix[i]):
            intensity = int(255 - (val / max_v) * 160)
            fill = (intensity, 230, 255)
            x0 = left + j * cell_w
            y0 = top + i * cell_h
            d.rectangle((x0, y0, x0 + cell_w, y0 + cell_h), fill=fill, outline="#CBD5E1")
            draw_text(d, (x0 + cell_w // 2, y0 + cell_h // 2), f"{val:.0%}" if val <= 1 else f"{val:.0f}", value_font, anchor="mm")
    img.save(path)


def save_line_chart(series: Dict[str, List[float]], x_labels: List[str], path: Path, title: str, unit: str = "") -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    w, h = 1200, 720
    img = Image.new("RGB", (w, h), "white")
    d = ImageDraw.Draw(img)
    title_font = get_font(34)
    label_font = get_font(18)
    draw_text(d, (w // 2, 34), title, title_font, anchor="ma")
    left, right, top, bottom = 110, 220, 110, 110
    plot_w = w - left - right
    plot_h = h - top - bottom
    all_vals = [v for vals in series.values() for v in vals]
    max_v = max(all_vals or [1]) or 1
    d.line((left, top, left, top + plot_h), fill="#333333", width=2)
    d.line((left, top + plot_h, left + plot_w, top + plot_h), fill="#333333", width=2)
    colors = ["#2563EB", "#059669", "#D97706", "#DC2626", "#7C3AED"]
    for idx, (name, vals) in enumerate(series.items()):
        pts = []
        for i, v in enumerate(vals):
            x = left + (plot_w * i / max(1, len(x_labels) - 1))
            y = top + plot_h - (v / max_v) * (plot_h - 24)
            pts.append((x, y))
        if len(pts) >= 2:
            d.line(pts, fill=colors[idx % len(colors)], width=4)
        for x, y in pts:
            d.ellipse((x - 5, y - 5, x + 5, y + 5), fill=colors[idx % len(colors)])
        legend_y = top + idx * 42
        d.line((w - right + 36, legend_y, w - right + 86, legend_y), fill=colors[idx % len(colors)], width=5)
        draw_text(d, (w - right + 98, legend_y - 12), name, label_font)
    for i, lab in enumerate(x_labels):
        x = left + (plot_w * i / max(1, len(x_labels) - 1))
        draw_text(d, (x, top + plot_h + 18), lab, label_font, anchor="ma")
    img.save(path)


def save_scatter_chart(points: List[Dict[str, Any]], path: Path, title: str, x_label: str = "X", y_label: str = "Y") -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    w, h = 1100, 720
    img = Image.new("RGB", (w, h), "white")
    d = ImageDraw.Draw(img)
    title_font = get_font(34)
    label_font = get_font(18)
    small_font = get_font(15)
    draw_text(d, (w // 2, 34), title, title_font, anchor="ma")
    left, right, top, bottom = 105, 70, 110, 115
    plot_w = w - left - right
    plot_h = h - top - bottom
    xs = [float(p.get("x", 0) or 0) for p in points]
    ys = [float(p.get("y", 0) or 0) for p in points]
    max_x = max(xs or [1]) or 1
    max_y = max(ys or [1]) or 1
    d.line((left, top, left, top + plot_h), fill="#333333", width=2)
    d.line((left, top + plot_h, left + plot_w, top + plot_h), fill="#333333", width=2)
    draw_text(d, (left + plot_w // 2, h - 44), x_label, label_font, anchor="ma")
    draw_text(d, (35, top + plot_h // 2), y_label, label_font, anchor="mm")
    for p in points:
        x = left + float(p.get("x", 0) or 0) / max_x * (plot_w - 30)
        y = top + plot_h - float(p.get("y", 0) or 0) / max_y * (plot_h - 30)
        color = p.get("color") or "#2563EB"
        d.ellipse((x - 7, y - 7, x + 7, y + 7), fill=color, outline="#1E293B")
        label = str(p.get("label", ""))
        if label:
            draw_text(d, (x + 10, y - 10), label, small_font)
    img.save(path)


def save_grouped_bar_chart(groups: Dict[str, Dict[str, float]], path: Path, title: str, unit: str = "") -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    group_labels = list(groups.keys())
    series_labels = sorted({k for vals in groups.values() for k in vals})
    w, h = 1280, 760
    img = Image.new("RGB", (w, h), "white")
    d = ImageDraw.Draw(img)
    title_font = get_font(34)
    label_font = get_font(17)
    small_font = get_font(15)
    draw_text(d, (w // 2, 34), title, title_font, anchor="ma")
    left, right, top, bottom = 105, 220, 115, 145
    plot_w = w - left - right
    plot_h = h - top - bottom
    max_v = max([v for vals in groups.values() for v in vals.values()] or [1]) or 1
    colors = ["#2563EB", "#059669", "#D97706", "#DC2626", "#7C3AED", "#0891B2"]
    d.line((left, top, left, top + plot_h), fill="#333333", width=2)
    d.line((left, top + plot_h, left + plot_w, top + plot_h), fill="#333333", width=2)
    slot_w = plot_w / max(1, len(group_labels))
    bar_w = max(14, min(42, slot_w / max(1, len(series_labels) + 1)))
    for i, group in enumerate(group_labels):
        base_x = left + i * slot_w + slot_w * 0.12
        for j, series in enumerate(series_labels):
            val = groups[group].get(series, 0)
            x0 = base_x + j * bar_w
            y1 = top + plot_h
            y0 = y1 - val / max_v * (plot_h - 28)
            d.rectangle((x0, y0, x0 + bar_w - 2, y1), fill=colors[j % len(colors)])
        draw_text(d, (left + i * slot_w + slot_w / 2, top + plot_h + 16), group, label_font, anchor="ma")
    for j, series in enumerate(series_labels):
        y = top + j * 38
        d.rectangle((w - right + 28, y, w - right + 54, y + 22), fill=colors[j % len(colors)])
        draw_text(d, (w - right + 66, y - 4), series, small_font)
    img.save(path)


def save_stacked_bar_chart(data: Dict[str, Dict[str, float]], path: Path, title: str) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    labels = list(data.keys())
    cats = sorted({c for vals in data.values() for c in vals})
    w, h = 1200, 760
    img = Image.new("RGB", (w, h), "white")
    d = ImageDraw.Draw(img)
    title_font = get_font(34)
    label_font = get_font(18)
    draw_text(d, (w // 2, 34), title, title_font, anchor="ma")
    left, right, top, bottom = 110, 260, 115, 145
    plot_w = w - left - right
    plot_h = h - top - bottom
    max_total = max([sum(vals.values()) for vals in data.values()] or [1]) or 1
    bar_w = max(44, int(plot_w / max(1, len(labels)) * 0.55))
    colors = ["#2563EB", "#059669", "#D97706", "#DC2626", "#7C3AED", "#0891B2", "#64748B"]
    d.line((left, top, left, top + plot_h), fill="#333333", width=2)
    d.line((left, top + plot_h, left + plot_w, top + plot_h), fill="#333333", width=2)
    for i, lab in enumerate(labels):
        x0 = left + (i + 0.5) * plot_w / max(1, len(labels)) - bar_w / 2
        y = top + plot_h
        for j, cat in enumerate(cats):
            val = data[lab].get(cat, 0)
            seg_h = int(val / max_total * (plot_h - 20))
            d.rectangle((x0, y - seg_h, x0 + bar_w, y), fill=colors[j % len(colors)], outline="white")
            y -= seg_h
        draw_text(d, (x0 + bar_w / 2, top + plot_h + 16), str(lab), label_font, anchor="ma")
    for j, cat in enumerate(cats):
        y = top + j * 40
        d.rectangle((w - right + 30, y, w - right + 58, y + 24), fill=colors[j % len(colors)])
        draw_text(d, (w - right + 72, y - 4), cat, label_font)
    img.save(path)


def create_docx_report(
    path: Path,
    title: str,
    sections: List[Dict[str, Any]],
    subtitle: str = "",
) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)
    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10.5)
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run(title)
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(31, 78, 121)
    if subtitle:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(subtitle)
        r.font.size = Pt(10)
        r.font.color.rgb = RGBColor(89, 89, 89)
    for sec in sections:
        heading = sec.get("heading")
        if heading:
            doc.add_heading(heading, level=1)
        for para in sec.get("paragraphs", []):
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(6)
            p.add_run(str(para))
        for bullets in sec.get("bullets", []):
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(str(bullets))
        for table_def in sec.get("tables", []):
            rows = table_def.get("rows", [])
            cols = table_def.get("columns", [])
            if rows and cols:
                if table_def.get("title"):
                    cap = doc.add_paragraph()
                    cap.add_run(table_def["title"]).bold = True
                table = doc.add_table(rows=1, cols=len(cols))
                table.style = "Table Grid"
                for i, col in enumerate(cols):
                    table.rows[0].cells[i].text = str(col)
                for row in rows:
                    cells = table.add_row().cells
                    for i, col in enumerate(cols):
                        cells[i].text = str(row.get(col, ""))
        for img_path in sec.get("images", []):
            img_path = Path(img_path)
            if img_path.exists():
                doc.add_picture(str(img_path), width=Inches(6.5))
                last = doc.paragraphs[-1]
                last.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.save(path)


def create_pdf_text(path: Path, title: str, sections: List[Dict[str, Any]], subtitle: str = "") -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    font_name = "Helvetica"
    for font_path in [
        "/System/Library/Fonts/Supplemental/Arial Unicode.ttf",
        "/System/Library/Fonts/STHeiti Medium.ttc",
        "/System/Library/Fonts/Hiragino Sans GB.ttc",
    ]:
        if Path(font_path).exists():
            try:
                pdfmetrics.registerFont(TTFont("EmbeddedCJK", font_path))
                font_name = "EmbeddedCJK"
                break
            except Exception:
                continue
    if font_name == "Helvetica":
        try:
            pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))
            font_name = "STSong-Light"
        except Exception:
            font_name = "Helvetica"
    c = canvas.Canvas(str(path), pagesize=A4)
    width, height = A4
    margin = 48
    y = height - margin
    max_text_width = width - margin * 2

    def wrap_to_width(text: str, size: int) -> List[str]:
        wrapped: List[str] = []
        current = ""
        for ch in str(text):
            if ch == "\n":
                wrapped.append(current)
                current = ""
                continue
            candidate = current + ch
            if not current or c.stringWidth(candidate, font_name, size) <= max_text_width:
                current = candidate
            else:
                wrapped.append(current.rstrip())
                current = ch.lstrip()
        wrapped.append(current)
        return wrapped or [""]

    def line(text: str, size: int = 10, leading: int = 15, bold: bool = False):
        nonlocal y
        c.setFont(font_name, size)
        for part in wrap_to_width(str(text), size):
            if y < margin + leading:
                c.showPage()
                y = height - margin
                c.setFont(font_name, size)
            c.drawString(margin, y, part)
            y -= leading

    c.setFont(font_name, 17)
    c.drawCentredString(width / 2, y, title)
    y -= 24
    if subtitle:
        c.setFont(font_name, 10)
        c.drawCentredString(width / 2, y, subtitle)
        y -= 24
    for sec in sections:
        if sec.get("heading"):
            y -= 6
            line(sec["heading"], size=13, leading=20)
        for para in sec.get("paragraphs", []):
            line(para, size=10, leading=16)
            y -= 3
        for bullet in sec.get("bullets", []):
            line("• " + str(bullet), size=10, leading=16)
        for table_def in sec.get("tables", []):
            if table_def.get("title"):
                line(table_def["title"], size=10, leading=16)
            cols = table_def.get("columns", [])
            rows = table_def.get("rows", [])[:12]
            if cols:
                line(" | ".join(map(str, cols)), size=8, leading=13)
                for row in rows:
                    line(" | ".join(str(row.get(col, "")) for col in cols), size=8, leading=13)
        for img_path in sec.get("images", []):
            img_path = Path(img_path)
            if img_path.exists():
                if y < 250:
                    c.showPage()
                    y = height - margin
                c.drawImage(str(img_path), margin, y - 220, width=width - margin * 2, height=210, preserveAspectRatio=True, anchor="n")
                y -= 230
    c.save()


def top_items(counter: Counter, n: int = 8) -> List[Dict[str, Any]]:
    total = sum(counter.values()) or 1
    return [{"item": k, "count": v, "share": round(v / total, 4)} for k, v in counter.most_common(n)]


def copy_alias(src: Path, alias: Path) -> None:
    alias.parent.mkdir(parents=True, exist_ok=True)
    if src.exists() and src.resolve() != alias.resolve():
        alias.write_bytes(src.read_bytes())

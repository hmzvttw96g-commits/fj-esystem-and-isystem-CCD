from __future__ import annotations

import re
from html import escape
from pathlib import Path
from typing import Dict, List

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.shared import Inches, Pt, RGBColor

from common import GENERATED_DIR, LOGS_DIR, PROCESSED_DIR, create_pdf_text, read_xlsx_rows, setup_logger, write_xlsx


logger = setup_logger("generate_two_prediction_papers", LOGS_DIR / "generate_two_prediction_papers.log")

FULL_SCORE = 100
STRUCTURE = [
    ("选择题", "一、选择题（本大题共15小题，每小题3分，共45分）"),
    ("填空题", "二、填空题（本大题共5小题，每小题3分，共15分）"),
    ("解答题", "三、解答题（本大题共4小题，每小题10分，共40分）"),
]

MATH_NS = "http://schemas.openxmlformats.org/officeDocument/2006/math"
MATH_RE = re.compile(
    "|".join(
        [
            r"\([xy0-9+\-][^)]*\)\^2[+\-]\([xy0-9+\-][^)]*\)\^2\s*=\s*[\d+\-*/×÷^().]+(?:=\d+)?",
            r"[A-Za-z]\([^)]*\)\s*=\s*\{[^{}]+\}",
            r"[A-Za-z]\([^)]*\)\s*=\s*[A-Za-z0-9_+\-*/×÷^().]+(?:\s*=\s*[A-Za-z0-9_+\-*/×÷^().]+)+",
            r"\|[^|]{1,12}\|\s*=",
            r"[-+]?\d+\s*<\s*[A-Za-z]\s*<=?\s*[-+]?\d+",
            r"[A-Za-z]\s*=\s*[-+]?\d+\s*<\s*[-+]?\d+",
            r"(?:sqrt|log|sin|cos|tan)_?[A-Za-z0-9]*\s*[A-Za-z0-9°()]*\s*=\s*[A-Za-z0-9_+\-*/×÷^().]+",
            r"[A-Za-z]\([^)]*\)\s*=\s*[-+]?\d+",
            r"[A-Za-z]\([^)]*\)\s*=\s*(?:（\s*）|[A-Za-z0-9_+\-*/×÷^().]+)",
            r"[A-Za-z]\s*=\s*\(\s*[-+]?\d+\s*,\s*[-+]?\d+\s*\)",
            r"[A-Z]\s*[∩∪]\s*[A-Z]\s*=",
            r"[\u4e00-\u9fa5A-Za-z0-9_]+(?:\([^)]+\))?\s*=\s*[A-Za-z0-9_+\-*/×÷^().]+(?:\s*(?:=|<=|>=|≤|≥|<|>)\s*[A-Za-z0-9_+\-*/×÷^().]+)*",
            r"[A-Za-z0-9_+\-*/×÷^().]+\s*(?:<=|>=|≤|≥|<|>|=)\s*(?:（\s*）|[A-Za-z0-9_+\-*/×÷^().]+)",
            r"[A-Z]\s*=\s*\{[^{}]+\}",
            r"\{[^{}]+\}",
            r"[A-Z]\s*[∩∪]\s*[A-Z]",
            r"[A-Za-z]\([A-Za-z0-9+\-*/×÷^_ ,，<>≤≥=]*\)",
            r"[\[\(]\s*[-+]?\d+\s*,\s*(?:\+?∞|[-+]?\d+)\s*[\]\)]",
            r"\(\s*[-+]?\d+\s*,\s*[-+]?\d+\s*\)",
            r"\|[^|]{1,12}\|",
            r"(?:sqrt|log|sin|cos|tan)_[A-Za-z0-9]+\s*[A-Za-z0-9]+",
            r"(?:sqrt|log|sin|cos|tan)\(?[A-Za-z0-9+\-*/×÷^_ ]*\)?°?",
            r"(?:\([^)]{1,30}\)|[A-Za-z0-9_]+(?:\^[+-]?\d+)?)(?:\s*(?:[+\-*/×÷]|×)\s*(?:\([^)]{1,30}\)|[A-Za-z0-9_]+(?:\^[+-]?\d+)?))*\s*(?:<=|>=|≤|≥|<|>|=)\s*[A-Za-z0-9_+\-*/×÷^().]+",
            r"[A-Za-z]+_[A-Za-z0-9]+",
            r"[A-Za-z]\d+",
            r"[A-Za-z0-9]+(?:\^[+-]?\d+)(?:[+\-*/×÷][A-Za-z0-9]+(?:\^[+-]?\d+)?)*",
        ]
    )
)


PAPERS: Dict[str, List[Dict[str, object]]] = {
    "A": [
        {
            "id": "1",
            "type": "选择题",
            "score": 3,
            "kp": "集合与逻辑",
            "detail": "集合交并补运算",
            "diff": "基础",
            "text": "设集合 A={2,4,6,8}，B={1,2,3,4,5}，则 A∩B=（  ）。\nA. {2,4}  B. {1,3,5}  C. {2,4,6,8}  D. {1,2,3,4,5,6,8}",
            "answer": "A",
            "solution": "A与B的公共元素是2和4，所以A∩B={2,4}。",
        },
        {
            "id": "2",
            "type": "选择题",
            "score": 3,
            "kp": "集合与逻辑",
            "detail": "区间与描述法",
            "diff": "基础",
            "text": "不等式 -1<=x<4 的解集用区间表示为（  ）。\nA. (-1,4)  B. [-1,4)  C. (-1,4]  D. [-1,4]",
            "answer": "B",
            "solution": "-1取到，4取不到，所以为[-1,4)。",
        },
        {
            "id": "3",
            "type": "选择题",
            "score": 3,
            "kp": "不等式",
            "detail": "一元一次不等式",
            "diff": "基础",
            "text": "不等式 4x-3>9 的解集是（  ）。\nA. x>3  B. x>=3  C. x<3  D. x<=3",
            "answer": "A",
            "solution": "4x-3>9，4x>12，所以x>3。",
        },
        {
            "id": "4",
            "type": "选择题",
            "score": 3,
            "kp": "函数",
            "detail": "函数定义域",
            "diff": "中等",
            "text": "函数 y=1/(x-2) 的定义域是（  ）。\nA. x为任意实数  B. x不等于2  C. x>2  D. x>=2",
            "answer": "B",
            "solution": "分母不能为0，所以x-2不等于0，即x不等于2。",
        },
        {
            "id": "5",
            "type": "选择题",
            "score": 3,
            "kp": "函数",
            "detail": "函数单调性",
            "diff": "中等",
            "text": "一次函数 y=-2x+5 在实数范围内是（  ）。\nA. 增函数  B. 减函数  C. 常函数  D. 不能判断",
            "answer": "B",
            "solution": "一次函数y=kx+b中，k=-2<0，所以函数为减函数。",
        },
        {
            "id": "6",
            "type": "选择题",
            "score": 3,
            "kp": "指数与对数",
            "detail": "对数运算",
            "diff": "中等",
            "text": "log_2 8 的值是（  ）。\nA. 2  B. 3  C. 4  D. 8",
            "answer": "B",
            "solution": "因为2^3=8，所以log_2 8=3。",
        },
        {
            "id": "7",
            "type": "选择题",
            "score": 3,
            "kp": "三角函数",
            "detail": "特殊角三角函数",
            "diff": "基础",
            "text": "cos30° 的值是（  ）。\nA. 1/2  B. sqrt(2)/2  C. sqrt(3)/2  D. 1",
            "answer": "C",
            "solution": "特殊角三角函数值：cos30°=sqrt(3)/2。",
        },
        {
            "id": "8",
            "type": "选择题",
            "score": 3,
            "kp": "数列",
            "detail": "等差数列",
            "diff": "基础",
            "text": "等差数列 1,4,7,... 的第5项是（  ）。\nA. 10  B. 12  C. 13  D. 16",
            "answer": "C",
            "solution": "公差为3，第5项a5=1+4×3=13。",
        },
        {
            "id": "9",
            "type": "选择题",
            "score": 3,
            "kp": "概率统计",
            "detail": "抽样与频率分布",
            "diff": "基础",
            "text": "为了了解某校1200名学生的视力情况，随机抽取120名学生调查，这120名学生称为（  ）。\nA. 总体  B. 个体  C. 样本  D. 样本容量",
            "answer": "C",
            "solution": "被抽取调查的120名学生构成样本。",
        },
        {
            "id": "10",
            "type": "选择题",
            "score": 3,
            "kp": "立体几何",
            "detail": "立体几何体积",
            "diff": "基础",
            "text": "长方体长、宽、高分别为5、4、3，则体积为（  ）。\nA. 12  B. 20  C. 60  D. 80",
            "answer": "C",
            "solution": "长方体体积=5×4×3=60。",
        },
        {
            "id": "11",
            "type": "选择题",
            "score": 3,
            "kp": "集合与逻辑",
            "detail": "充分必要条件",
            "diff": "中等",
            "text": "“x>2”是“x>0”的（  ）。\nA. 充分不必要条件  B. 必要不充分条件  C. 充要条件  D. 既不充分也不必要条件",
            "answer": "A",
            "solution": "x>2一定推出x>0；但x>0不一定推出x>2，所以是充分不必要条件。",
        },
        {
            "id": "12",
            "type": "选择题",
            "score": 3,
            "kp": "直线与圆",
            "detail": "直线位置关系",
            "diff": "中等",
            "text": "直线 y=2x+1 与直线 y=2x-3 的位置关系是（  ）。\nA. 相交但不垂直  B. 平行  C. 垂直  D. 重合",
            "answer": "B",
            "solution": "两条直线斜率相同、截距不同，所以平行。",
        },
        {
            "id": "13",
            "type": "选择题",
            "score": 3,
            "kp": "向量",
            "detail": "向量坐标运算",
            "diff": "中等",
            "text": "已知向量 a=(2,3)，b=(1,-2)，则 a-b=（  ）。\nA. (1,5)  B. (3,1)  C. (1,1)  D. (-1,-5)",
            "answer": "A",
            "solution": "a-b=(2-1,3-(-2))=(1,5)。",
        },
        {
            "id": "14",
            "type": "选择题",
            "score": 3,
            "kp": "直线与圆",
            "detail": "圆的标准方程",
            "diff": "中等",
            "text": "圆 (x-3)^2+(y+1)^2=4 的圆心是（  ）。\nA. (3,-1)  B. (-3,1)  C. (3,1)  D. (-3,-1)",
            "answer": "A",
            "solution": "标准方程(x-a)^2+(y-b)^2=r^2，圆心为(a,b)=(3,-1)。",
        },
        {
            "id": "15",
            "type": "选择题",
            "score": 3,
            "kp": "概率统计",
            "detail": "计数原理",
            "diff": "基础",
            "text": "某食堂有3种主食、4种菜品，各选1种，共有（  ）种搭配。\nA. 7  B. 12  C. 24  D. 34",
            "answer": "B",
            "solution": "分步计数：3×4=12种。",
        },
        {
            "id": "16",
            "type": "填空题",
            "score": 3,
            "kp": "不等式",
            "detail": "一元二次不等式",
            "diff": "中等",
            "text": "不等式 (x-2)(x+1)>=0 的解集为__________。",
            "answer": "x<=-1或x>=2",
            "solution": "两根为-1和2，二次式开口向上，大于等于0取两侧，所以x<=-1或x>=2。",
        },
        {
            "id": "17",
            "type": "填空题",
            "score": 3,
            "kp": "函数",
            "detail": "函数求值",
            "diff": "基础",
            "text": "已知 f(x)=x^2-2x，则 f(3)=__________。",
            "answer": "3",
            "solution": "f(3)=3^2-2×3=9-6=3。",
        },
        {
            "id": "18",
            "type": "填空题",
            "score": 3,
            "kp": "数列",
            "detail": "等比数列",
            "diff": "中等",
            "text": "等比数列首项为2，公比为-3，则第3项为__________。",
            "answer": "18",
            "solution": "a3=2×(-3)^2=18。",
        },
        {
            "id": "19",
            "type": "填空题",
            "score": 3,
            "kp": "概率统计",
            "detail": "概率基础",
            "diff": "基础",
            "text": "从数字1、2、3、4中任取一个，取到偶数的概率为__________。",
            "answer": "1/2",
            "solution": "偶数有2、4两个，共4个数，所以概率为2/4=1/2。",
        },
        {
            "id": "20",
            "type": "填空题",
            "score": 3,
            "kp": "直线与圆",
            "detail": "直线斜率",
            "diff": "基础",
            "text": "经过点(0,1)和(2,5)的直线斜率为__________。",
            "answer": "2",
            "solution": "k=(5-1)/(2-0)=2。",
        },
        {
            "id": "21",
            "type": "解答题",
            "score": 10,
            "kp": "集合与逻辑",
            "detail": "集合交并补运算",
            "diff": "中等",
            "text": "已知全集 U={1,2,3,4,5,6,7,8,9}，集合 A={1,3,5,7,9}，B={3,4,5,6}。\n(1) 求 A∩B；\n(2) 求 A∪B；\n(3) 求集合 B 在全集 U 中的补集。",
            "answer": "(1){3,5}；(2){1,3,4,5,6,7,9}；(3){1,2,7,8,9}",
            "solution": "A∩B={3,5}。A∪B={1,3,4,5,6,7,9}。U中不属于B的元素为{1,2,7,8,9}。",
        },
        {
            "id": "22",
            "type": "解答题",
            "score": 10,
            "kp": "函数",
            "detail": "一次函数模型",
            "diff": "较难",
            "text": "某班购买练习册，每本6元，另需一次性包装费30元。设购买x本的总费用为y元。\n(1) 写出y关于x的函数关系式；\n(2) 购买40本需要多少元；\n(3) 若预算不超过330元，最多可购买多少本？",
            "answer": "(1)y=6x+30；(2)270元；(3)50本",
            "solution": "总费用为y=6x+30。x=40时，y=6×40+30=270。6x+30<=330，得x<=50，最多50本。",
        },
        {
            "id": "23",
            "type": "解答题",
            "score": 10,
            "kp": "直线与圆",
            "detail": "直线与圆综合",
            "diff": "较难",
            "text": "已知直线 l 经过点 A(1,1)、B(4,7)，圆 C 的方程为 (x-1)^2+(y-2)^2=25。\n(1) 求直线 l 的斜率；\n(2) 写出圆 C 的圆心和半径；\n(3) 判断点 B 是否在圆 C 上。",
            "answer": "(1)2；(2)圆心(1,2)，半径5；(3)不在圆上",
            "solution": "k=(7-1)/(4-1)=2。圆心为(1,2)，半径为5。点B代入得(4-1)^2+(7-2)^2=9+25=34，不等于25，所以不在圆上。",
        },
        {
            "id": "24",
            "type": "解答题",
            "score": 10,
            "kp": "数列",
            "detail": "等差数列",
            "diff": "较难",
            "text": "某实训小组第一天完成零件12个，以后每天比前一天多完成3个。\n(1) 第5天完成多少个；\n(2) 写出第n天完成零件数a_n；\n(3) 前5天共完成多少个？",
            "answer": "(1)24个；(2)a_n=12+3(n-1)；(3)90个",
            "solution": "这是等差数列，a1=12，d=3。a5=12+4×3=24。a_n=12+3(n-1)。前5天和S5=(12+24)×5/2=90。",
        },
    ],
    "B": [
        {
            "id": "1",
            "type": "选择题",
            "score": 3,
            "kp": "集合与逻辑",
            "detail": "集合表示与元素关系",
            "diff": "基础",
            "text": "下列元素属于集合 A={x|x为小于5的正整数} 的是（  ）。\nA. 0  B. 3  C. 5  D. -1",
            "answer": "B",
            "solution": "小于5的正整数为1、2、3、4，所以3属于集合A。",
        },
        {
            "id": "2",
            "type": "选择题",
            "score": 3,
            "kp": "集合与逻辑",
            "detail": "集合交并补运算",
            "diff": "基础",
            "text": "已知全集 U={1,2,3,4,5}，A={1,3,5}，则 A 在 U 中的补集是（  ）。\nA. {1,3,5}  B. {2,4}  C. {1,2,3,4,5}  D. 空集",
            "answer": "B",
            "solution": "U中不属于A的元素是2和4。",
        },
        {
            "id": "3",
            "type": "选择题",
            "score": 3,
            "kp": "不等式",
            "detail": "绝对值不等式",
            "diff": "中等",
            "text": "不等式 |x|<3 的解集是（  ）。\nA. x<3  B. x>-3  C. -3<x<3  D. x<-3或x>3",
            "answer": "C",
            "solution": "|x|<3表示x到0的距离小于3，所以-3<x<3。",
        },
        {
            "id": "4",
            "type": "选择题",
            "score": 3,
            "kp": "函数",
            "detail": "函数奇偶性",
            "diff": "中等",
            "text": "下列函数中是奇函数的是（  ）。\nA. y=x  B. y=x^2  C. y=x^2+1  D. y=2",
            "answer": "A",
            "solution": "y=x满足f(-x)=-f(x)，是奇函数。",
        },
        {
            "id": "5",
            "type": "选择题",
            "score": 3,
            "kp": "函数",
            "detail": "函数求值",
            "diff": "基础",
            "text": "已知 f(x)=3x-4，则 f(5)=（  ）。\nA. 9  B. 11  C. 15  D. 19",
            "answer": "B",
            "solution": "f(5)=3×5-4=11。",
        },
        {
            "id": "6",
            "type": "选择题",
            "score": 3,
            "kp": "指数与对数",
            "detail": "指数幂运算",
            "diff": "基础",
            "text": "2^3×2^2=（  ）。\nA. 2^5  B. 2^6  C. 4^5  D. 4^6",
            "answer": "A",
            "solution": "同底数幂相乘，指数相加，2^3×2^2=2^5。",
        },
        {
            "id": "7",
            "type": "选择题",
            "score": 3,
            "kp": "三角函数",
            "detail": "三角函数符号与终边",
            "diff": "中等",
            "text": "角120°的终边在第（  ）象限。\nA. 一  B. 二  C. 三  D. 四",
            "answer": "B",
            "solution": "120°在90°到180°之间，终边在第二象限。",
        },
        {
            "id": "8",
            "type": "选择题",
            "score": 3,
            "kp": "数列",
            "detail": "等比数列",
            "diff": "基础",
            "text": "等比数列 2,6,18,... 的公比是（  ）。\nA. 2  B. 3  C. 6  D. 18",
            "answer": "B",
            "solution": "公比为6/2=3。",
        },
        {
            "id": "9",
            "type": "选择题",
            "score": 3,
            "kp": "概率统计",
            "detail": "概率基础",
            "diff": "基础",
            "text": "抛掷一枚均匀硬币一次，出现正面的概率为（  ）。\nA. 0  B. 1/2  C. 1  D. 2",
            "answer": "B",
            "solution": "硬币正反两面等可能，出现正面的概率为1/2。",
        },
        {
            "id": "10",
            "type": "选择题",
            "score": 3,
            "kp": "立体几何",
            "detail": "立体几何视图与位置",
            "diff": "中等",
            "text": "圆柱的侧面展开图通常是（  ）。\nA. 三角形  B. 圆  C. 矩形  D. 梯形",
            "answer": "C",
            "solution": "圆柱侧面沿母线展开后是矩形。",
        },
        {
            "id": "11",
            "type": "选择题",
            "score": 3,
            "kp": "集合与逻辑",
            "detail": "充分必要条件",
            "diff": "中等",
            "text": "“x=0”是“xy=0”的（  ）。\nA. 充分不必要条件  B. 必要不充分条件  C. 充要条件  D. 既不充分也不必要条件",
            "answer": "A",
            "solution": "x=0一定有xy=0；但xy=0也可能是y=0，所以不是必要条件。",
        },
        {
            "id": "12",
            "type": "选择题",
            "score": 3,
            "kp": "向量",
            "detail": "向量坐标运算",
            "diff": "中等",
            "text": "已知向量 a=(3,4)，则 |a|=（  ）。\nA. 3  B. 4  C. 5  D. 7",
            "answer": "C",
            "solution": "|a|=sqrt(3^2+4^2)=5。",
        },
        {
            "id": "13",
            "type": "选择题",
            "score": 3,
            "kp": "直线与圆",
            "detail": "直线斜率",
            "diff": "基础",
            "text": "经过点(1,2)、(3,8)的直线斜率为（  ）。\nA. 2  B. 3  C. 4  D. 6",
            "answer": "B",
            "solution": "k=(8-2)/(3-1)=6/2=3。",
        },
        {
            "id": "14",
            "type": "选择题",
            "score": 3,
            "kp": "直线与圆",
            "detail": "圆的标准方程",
            "diff": "中等",
            "text": "以原点为圆心、半径为6的圆的方程是（  ）。\nA. x^2+y^2=6  B. x^2+y^2=12  C. x^2+y^2=36  D. (x-6)^2+y^2=36",
            "answer": "C",
            "solution": "圆心在原点、半径为6，则方程为x^2+y^2=36。",
        },
        {
            "id": "15",
            "type": "选择题",
            "score": 3,
            "kp": "概率统计",
            "detail": "计数原理",
            "diff": "基础",
            "text": "从甲、乙、丙3人中选1人担任组长，再从剩下2人中选1人担任记录员，共有（  ）种选法。\nA. 3  B. 5  C. 6  D. 9",
            "answer": "C",
            "solution": "先选组长3种，再选记录员2种，共3×2=6种。",
        },
        {
            "id": "16",
            "type": "填空题",
            "score": 3,
            "kp": "集合与逻辑",
            "detail": "区间与描述法",
            "diff": "基础",
            "text": "区间 (0,5] 用不等式表示为__________。",
            "answer": "0<x<=5",
            "solution": "左端点0不取，右端点5取到，所以0<x<=5。",
        },
        {
            "id": "17",
            "type": "填空题",
            "score": 3,
            "kp": "不等式",
            "detail": "一元一次不等式",
            "diff": "基础",
            "text": "不等式 5-2x<=1 的解集为__________。",
            "answer": "x>=2",
            "solution": "5-2x<=1，-2x<=-4，两边除以-2不等号变向，x>=2。",
        },
        {
            "id": "18",
            "type": "填空题",
            "score": 3,
            "kp": "函数",
            "detail": "函数定义域",
            "diff": "中等",
            "text": "函数 y=sqrt(2x-4) 的定义域为__________。",
            "answer": "x>=2",
            "solution": "根号内2x-4>=0，得x>=2。",
        },
        {
            "id": "19",
            "type": "填空题",
            "score": 3,
            "kp": "数列",
            "detail": "等差数列",
            "diff": "基础",
            "text": "等差数列首项为5，公差为4，则第6项为__________。",
            "answer": "25",
            "solution": "a6=5+5×4=25。",
        },
        {
            "id": "20",
            "type": "填空题",
            "score": 3,
            "kp": "直线与圆",
            "detail": "圆的标准方程",
            "diff": "中等",
            "text": "圆 (x-1)^2+(y+3)^2=9 的半径为__________。",
            "answer": "3",
            "solution": "r^2=9，所以半径r=3。",
        },
        {
            "id": "21",
            "type": "解答题",
            "score": 10,
            "kp": "不等式",
            "detail": "一元二次不等式",
            "diff": "较难",
            "text": "解下列不等式，并用区间表示结果：\n(1) 3x+1>=10；\n(2) (x-1)(x-5)<0；\n(3) 求两个解集的交集。",
            "answer": "(1)[3,+∞)；(2)(1,5)；(3)[3,5)",
            "solution": "(1)3x+1>=10，得x>=3。(2)两根为1和5，二次式小于0取中间，得1<x<5。(3)交集为[3,5)。",
        },
        {
            "id": "22",
            "type": "解答题",
            "score": 10,
            "kp": "函数",
            "detail": "分段函数",
            "diff": "较难",
            "text": "已知 f(x)={ x^2, x<2；2x+1, x>=2 }。\n(1) 求 f(1) 和 f(3)；\n(2) 若 x>=2 且 f(x)=9，求x；\n(3) 判断点(2,5)是否在函数图像上。",
            "answer": "(1)1，7；(2)4；(3)在",
            "solution": "x=1<2，f(1)=1^2=1。x=3>=2，f(3)=2×3+1=7。x>=2时2x+1=9，x=4。f(2)=2×2+1=5，所以点(2,5)在图像上。",
        },
        {
            "id": "23",
            "type": "解答题",
            "score": 10,
            "kp": "向量",
            "detail": "向量坐标运算",
            "diff": "较难",
            "text": "已知 A(0,0)，B(4,3)，C(4,0)。\n(1) 写出向量 AB 的坐标；\n(2) 求 |AB|；\n(3) 求三角形 ABC 的面积。",
            "answer": "(1)(4,3)；(2)5；(3)6",
            "solution": "AB=(4,3)。|AB|=sqrt(4^2+3^2)=5。AC=4，点B到x轴距离为3，面积=1/2×4×3=6。",
        },
        {
            "id": "24",
            "type": "解答题",
            "score": 10,
            "kp": "应用与建模",
            "detail": "应用函数建模",
            "diff": "较难",
            "text": "某校租车参加技能展示，固定服务费120元，每行驶1千米收费5元。设行驶x千米的总费用为y元。\n(1) 写出y关于x的函数关系式；\n(2) 行驶60千米需付多少元；\n(3) 若总费用不超过500元，最多可行驶多少千米？",
            "answer": "(1)y=5x+120；(2)420元；(3)76千米",
            "solution": "费用关系为y=5x+120。x=60时，y=300+120=420。5x+120<=500，得x<=76，最多76千米。",
        },
    ],
}


def load_lookup(path: Path, sheet: str, key: str) -> Dict[str, Dict[str, object]]:
    if not path.exists():
        return {}
    return {str(r.get(key)): r for r in read_xlsx_rows(path, sheet)}


def add_title(doc: Document, title: str, subtitle: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(18)
    r.font.color.rgb = RGBColor(31, 78, 121)
    sp = doc.add_paragraph()
    sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sp.add_run(subtitle)
    sr.font.size = Pt(10)
    sr.font.color.rgb = RGBColor(89, 89, 89)


def add_equation(paragraph, formula: str) -> None:
    formula = formula.strip()
    if not formula:
        return
    xml = (
        f'<m:oMath xmlns:m="{MATH_NS}">'
        f"<m:r><m:t xml:space=\"preserve\">{escape(formula)}</m:t></m:r>"
        f"</m:oMath>"
    )
    paragraph._p.append(parse_xml(xml))


def add_mixed_text(paragraph, text: str, bold: bool = False) -> None:
    pos = 0
    for match in MATH_RE.finditer(text):
        start, end = match.span()
        if start > pos:
            run = paragraph.add_run(text[pos:start])
            run.bold = bold
        add_equation(paragraph, match.group(0))
        pos = end
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        run.bold = bold


def add_math_paragraph(doc: Document, text: str, style: str | None = None):
    p = doc.add_paragraph(style=style)
    add_mixed_text(p, text)
    return p


def add_question_text(doc: Document, text: str) -> None:
    for line in str(text).splitlines():
        add_math_paragraph(doc, line)


def build_docx(path: Path, paper_name: str, questions: List[Dict[str, object]], answers: bool = False) -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)
    doc.styles["Normal"].font.size = Pt(10.5)
    suffix = " 答案解析" if answers else ""
    add_title(doc, f"福建省中职学业水平测试公共基础数学{paper_name}{suffix}", f"预测模拟用途｜满分{FULL_SCORE}分｜15选择+5填空+4解答")
    if not answers:
        doc.add_paragraph("注意事项：本卷只含公共基础综合卷中的数学部分。请在答题前核对题号，选择题每题只有一个正确答案。")
    for qtype, intro in STRUCTURE:
        doc.add_heading(intro, level=1)
        for q in [x for x in questions if x["type"] == qtype]:
            if answers:
                p = doc.add_paragraph()
                p.add_run(f"{q['id']}. 答案：").bold = True
                add_mixed_text(p, str(q["answer"]))
                p = doc.add_paragraph()
                p.add_run("解析：").bold = True
                add_mixed_text(p, str(q["solution"]))
            else:
                add_question_text(doc, f"{q['id']}. {q['text']}")
    doc.save(path)


def pdf_sections(questions: List[Dict[str, object]], answers: bool = False):
    sections = []
    for qtype, _intro in STRUCTURE:
        paras = []
        for q in [x for x in questions if x["type"] == qtype]:
            paras.append(f"{q['id']}. 答案：{q['answer']}。解析：{q['solution']}" if answers else f"{q['id']}. {q['text']}")
        sections.append({"heading": qtype, "paragraphs": paras})
    return sections


def validate_paper(name: str, questions: List[Dict[str, object]]) -> None:
    counts = {qtype: sum(1 for q in questions if q["type"] == qtype) for qtype, _ in STRUCTURE}
    score = sum(int(q["score"]) for q in questions)
    if counts != {"选择题": 15, "填空题": 5, "解答题": 4} or score != FULL_SCORE:
        raise ValueError(f"{name} structure invalid: counts={counts}, score={score}")


def main() -> None:
    logger.info("Start generating two prediction papers")
    ols_lookup = load_lookup(PROCESSED_DIR / "detailed_knowledge_ols_analysis.xlsx", "ols_weighted_blueprint", "detailed_knowledge_point")
    trend_lookup = load_lookup(PROCESSED_DIR / "future_exam_trend_prediction.xlsx", "trend_prediction", "knowledge_point")
    all_basis: List[Dict[str, object]] = []
    subtitle = f"预测模拟用途｜满分{FULL_SCORE}分｜15选择+5填空+4解答"

    for key, questions in PAPERS.items():
        validate_paper(key, questions)
        paper_name = f"押题卷{key}"
        paper_docx = GENERATED_DIR / f"福建省中职学业水平测试公共基础数学{paper_name}.docx"
        answer_docx = GENERATED_DIR / f"福建省中职学业水平测试公共基础数学{paper_name}_答案解析.docx"
        build_docx(paper_docx, paper_name, questions, answers=False)
        build_docx(answer_docx, paper_name, questions, answers=True)
        create_pdf_text(GENERATED_DIR / f"福建省中职学业水平测试公共基础数学{paper_name}.pdf", f"福建省中职学业水平测试公共基础数学{paper_name}", pdf_sections(questions, False), subtitle)
        create_pdf_text(GENERATED_DIR / f"福建省中职学业水平测试公共基础数学{paper_name}_答案解析.pdf", f"福建省中职学业水平测试公共基础数学{paper_name} 答案解析", pdf_sections(questions, True), subtitle)

        for q in questions:
            ols = ols_lookup.get(str(q["detail"]), {})
            trend = trend_lookup.get(str(q["kp"]), {})
            basis = "OLS细知识点权重"
            if not ols:
                basis = "考纲覆盖与大类趋势补位"
            all_basis.append(
                {
                    "paper": paper_name,
                    "question_id": q["id"],
                    "question_type": q["type"],
                    "score": q["score"],
                    "knowledge_point": q["kp"],
                    "detailed_knowledge_point": q["detail"],
                    "difficulty": q["diff"],
                    "ols_normalized_weight": ols.get("normalized_weight", ""),
                    "ols_recommended_questions_in_24": ols.get("recommended_questions_in_24", ""),
                    "historical_real_share": ols.get("historical_real_share", ""),
                    "xiamen_reference_share": ols.get("xiamen_reference_share", ""),
                    "prediction_score": trend.get("prediction_score", ""),
                    "recommended_coverage": trend.get("recommended_coverage", ""),
                    "design_basis": f"{basis}；同时满足15选、5填、4解的结构约束；题面原创，覆盖真题高频、厦门卷信号和近两年考纲边界。",
                    "answer": q["answer"],
                }
            )
        logger.info("Generated %s", paper_name)

    write_xlsx({"blueprint": all_basis}, GENERATED_DIR / "两套押题卷命题依据表.xlsx")
    logger.info("Output two prediction papers and blueprint")


if __name__ == "__main__":
    main()

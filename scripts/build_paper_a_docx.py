# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BODY="Songti SC"; HEAD="Heiti SC"
doc=Document()
sec=doc.sections[0]
sec.page_width=Inches(8.27); sec.page_height=Inches(11.69)   # A4
for m in ("top_margin","bottom_margin"): setattr(sec,m,Inches(1.0))
for m in ("left_margin","right_margin"): setattr(sec,m,Inches(1.1))

def setfont(run, name=BODY, size=10.5, bold=False, color=None):
    run.font.name=name; run.font.size=Pt(size); run.bold=bold
    if color: run.font.color.rgb=RGBColor(*color)
    rpr=run._element.get_or_add_rPr(); rf=rpr.find(qn('w:rFonts'))
    if rf is None: rf=OxmlElement('w:rFonts'); rpr.append(rf)
    rf.set(qn('w:eastAsia'),name); rf.set(qn('w:ascii'),name); rf.set(qn('w:hAnsi'),name)

def para(text="", size=10.5, bold=False, align=None, before=2, after=4, color=None, name=BODY):
    p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(before); p.paragraph_format.space_after=Pt(after)
    p.paragraph_format.line_spacing=1.4
    if align: p.alignment=align
    if text: setfont(p.add_run(text), name=name, size=size, bold=bold, color=color)
    return p

def heading(text, lvl=1):
    sizes={1:15,2:12.5,3:11.5}
    p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(12 if lvl==1 else 8); p.paragraph_format.space_after=Pt(5)
    p.paragraph_format.keep_with_next=True
    setfont(p.add_run(text), name=HEAD, size=sizes[lvl], bold=True, color=(0x1F,0x3A,0x4A))
    return p

def shade(cell, hexc):
    tcPr=cell._tc.get_or_add_tcPr(); sh=OxmlElement('w:shd')
    sh.set(qn('w:val'),'clear'); sh.set(qn('w:fill'),hexc); tcPr.append(sh)

def table(headers, rows, widths, header_fill="1F3A4A", header_color=(255,255,255), fontsize=9.5, zebra=True):
    t=doc.add_table(rows=1, cols=len(headers)); t.alignment=WD_TABLE_ALIGNMENT.CENTER
    t.style='Table Grid'
    for j,h in enumerate(headers):
        c=t.rows[0].cells[j]; c.width=Inches(widths[j]); shade(c,header_fill)
        c.paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER
        setfont(c.paragraphs[0].add_run(h), name=HEAD, size=fontsize, bold=True, color=header_color)
    for ri,row in enumerate(rows):
        cells=t.add_row().cells
        for j,val in enumerate(row):
            cells[j].width=Inches(widths[j])
            if zebra and ri%2==1: shade(cells[j],"F2F5F7")
            pp=cells[j].paragraphs[0]; pp.alignment=WD_ALIGN_PARAGRAPH.CENTER if j>0 else WD_ALIGN_PARAGRAPH.LEFT
            setfont(pp.add_run(str(val)), size=fontsize)
    return t

def figure(path, width, caption):
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before=Pt(6)
    p.add_run().add_picture(path, width=Inches(width))
    cap=doc.add_paragraph(); cap.alignment=WD_ALIGN_PARAGRAPH.CENTER; cap.paragraph_format.space_after=Pt(8)
    setfont(cap.add_run(caption), name=BODY, size=9, color=(0x66,0x66,0x66))

LINKS4=["供给规模","供给质量","产业承接","产业转化"]
def worked_table(city, v):
    """任意城市:由四环节标准化值 v 自动生成障碍度手算表(v→偏离度→份额%)。"""
    dev=[round(1-x,3) for x in v]; tot=round(sum(dev),3)
    sh=[round(d/tot*100,1) if tot>0 else 0.0 for d in dev]
    rows=[[LINKS4[i], f"{v[i]:.3f}", f"1−{v[i]:.3f} = {dev[i]:.3f}", f"{dev[i]:.3f} ÷ {tot:.3f} = {sh[i]:.1f}%"] for i in range(4)]
    rows.append(["合计","—",f"{tot:.3f}（总偏离）","100%"])
    table(["环节","①标准化值 v","②偏离度 = 1−v","③障碍度% = 偏离÷总偏离"],
          rows, [1.0,1.5,2.0,2.7], fontsize=9)

# ===== 标题 =====
ti=doc.add_paragraph(); ti.alignment=WD_ALIGN_PARAGRAPH.CENTER
setfont(ti.add_run("福建省 AI 人才供需匹配的教育—产业耦合协调测度"), name=HEAD, size=18, bold=True, color=(0x1F,0x3A,0x4A))
st=doc.add_paragraph(); st.alignment=WD_ALIGN_PARAGRAPH.CENTER; st.paragraph_format.space_after=Pt(10)
setfont(st.add_run("数据与检验"), name=HEAD, size=13, bold=True, color=(0x37,0x8A,0xDD))
para("摘要：本文以四环节功能链测度福建 9 设区市 2019—2024 年教育（E）与产业（I）的耦合协调度（CCD），"
     "并以障碍度模型诊断失调环节。结果显示福州、厦门优质协调（D>0.9）、其余 7 市长期失调，呈“双核—腹地”"
     "二元结构；福州“教育领先”、厦门“产业领先”构成镜像。本文详述障碍度的数据处理、方法与检验，"
     "并揭示 min-max 标准化对腹地城市障碍分解的压缩效应及其修正。", size=9.5, before=2, after=8)

# ===== 一 =====
heading("一、数据来源与口径", 1)
para("教育—产业耦合在“四环节功能链”上测度：教育系统（E）含供给规模与供给质量，产业系统（I）含产业承接"
     "与产业转化。各环节指标、来源与标准化方式如下。")
table(["系统","环节","指标","数据来源","标准化"],
 [["E","供给规模","AI 专业群本科招生计划数","福建省普通高校招生计划册（逐年解析）","人均化"],
  ["E","供给质量","国家级/省级一流本科专业建设点","教育部/省教育厅公开名单（逐条核验）","容量（非人均）"],
  ["I","产业承接","信息传输软件业城镇单位从业人员（门类I）","各市统计年鉴（官方核验）","人均化"],
  ["I","产业转化","企业 AI/数字技术专利申请量","智慧芽专利库（IPC 检索）","人均化"]],
 [0.7,1.0,2.4,2.3,0.9])
para("三口径严格嵌套（C⊂B⊂X）：C 为 AI 直接命名专业（人工智能、智能科学与技术、数据科学与大数据技术），"
     "B 加入计算机科学与技术、软件工程等基础专业群（主口径），X 进一步纳入电子信息、自动化、智能制造等交叉群；"
     "专利端按 IPC 前缀对应嵌套。人均化分母统一为全市·年末·常住人口，使指标度量“人均强度”而非城市规模。", before=4)

# ===== 二 =====
heading("二、数据处理流程", 1)
steps=[("抽取与归一","招生册按版式族自动解析（2019 Word/2020 两栏/2021 单栏左移/2022—2024 标准），校名与专业 join 口径配置；专利 75 396 件去重后剔高校/个人、IPC 分类、第一申请人地址归市，入账 35 169 件。"),
 ("质量闸门","完备性核查（零命中生成《真零核查单》）、年际跳变、加总锚（Σ九市≈全省）。"),
 ("标准化","全样本固定基准 min-max（按 C/B/X 分别做）；min/max 锚单元须数值审计签字。"),
 ("缺失处理","遵循“缺失不填 0、不插值、不外推”；可回推者仅限“全省锚可得+单年”（人口 2019 七普断点回推、漳州 2024 产业承接定锚回推；莆田 2019 代理值剔除）。"),
 ("测度","耦合协调度 D=√(C·T)，障碍度分解至四环节定位主障碍。")]
for i,(h,t) in enumerate(steps,1):
    p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(3); p.paragraph_format.line_spacing=1.4
    setfont(p.add_run(f"{i}. {h}："), bold=True, size=10); setfont(p.add_run(t), size=10)

# ===== 三 =====
heading("三、测度方法", 1)
para("耦合协调度：D=√(C·T)，其中耦合度 C=2√(E·I)/(E+I) 衡量 E、I 的协同程度，综合发展水平 T=(E+I)/2 衡量整体高度；"
     "D∈[0,1]，任一指数为 0 标“不可计算”，不参与等级与路径分类。")
para("障碍度模型：障碍度份额 = 因子贡献度 × 指标偏离度（归一化）。设环节 k 的标准化值为 vₖ、理想值为 1，"
     "则其偏离度为 (1−vₖ)、贡献度等权 1/4，障碍度份额 = (1−vₖ)/Σⱼ(1−vⱼ)。份额最大的环节即“主障碍”。", before=3)

# ===== 四 =====
heading("四、测量结果", 1)
heading("4.1  耦合协调度水平与时序", 2)
figure("/tmp/fig1_heatmap.png", 6.0, "图1  福建9市教育—产业耦合协调度热力图（B口径，2019—2024）。颜色越深协调度越高。")
para("核心发现一（双核—腹地分化）：福州、厦门 2024 年 D 均超 0.9（优质协调）且六年持续上升；其余 7 市长期处于 "
     "0.16—0.36 的失调区间。福建 AI 人才供需匹配呈鲜明的“双核协调、腹地失调”二元结构，且差距未随时间收敛。", before=4)

heading("4.2  E—I 结构与城市类型", 2)
figure("/tmp/fig2_scatter.png", 4.7, "图2  E—I 结构与双核镜像（B口径，2024）。对角线为供需均衡线。")
table(["市","E指数","I指数","D","城市类型","主障碍"],
 [["福州","1.000","0.769","0.937","教育领先·产业滞后","产业转化"],
  ["厦门","0.713","0.977","0.914","产业领先·教育滞后","供给规模"],
  ["泉州","0.258","0.067","0.363","双低失调","产业承接"],
  ["龙岩","0.083","0.107","0.306","双低失调","（弱信号）"],
  ["宁德","0.005","0.137","0.159","双低失调","供给质量"],
  ["莆田","0.081","0.044","0.244","双低失调","供给质量"],
  ["漳州","0.126","0.033","0.254","双低失调","产业承接"],
  ["三明","0.071","0.060","0.255","双低失调","（弱信号）"],
  ["南平","0.040","0.040","0.201","双低失调","（弱信号）"]],
 [0.7,1.0,1.0,0.9,2.6,1.6])
para("核心发现二（双核镜像）：福州与厦门同为高协调而机理相反。福州教育供给封顶（E=1.0）、产业转化滞后，落在均衡线"
     "“教育领先”一侧；厦门产业指数近满（I=0.977）、供给规模滞后，落在“产业领先”一侧。二者构成一对"
     "“高协调、反向失衡”的镜像——前者待产业承接其人才产出，后者待教育扩张其供给规模。", before=4)

# ===== 五 障碍详述 =====
heading("五、环节分析与障碍度诊断（详述）", 1)
heading("5.1  障碍度的数据处理：四环节标准化值的构造", 2)
para("障碍度直接建立在四环节的标准化值 vₖ 之上，其数据处理为两步：①人均化——供给规模、产业承接、产业转化三环节"
     "除以城市常住人口（供给质量为项目点计数，不人均化）；②全样本 min-max 标准化——按 C/B/X 口径分别，"
     "对“城市×年份”全样本取固定基准 min-max，将各环节映射至 [0,1]。下表为 B 口径 2024 年的标准化值与极差"
     "（四环节最大—最小），它是后续障碍分解可信度的关键。")
table(["市","供给规模","供给质量","产业承接","产业转化","极差","障碍可辨性"],
 [["福州","1.000","1.000","0.858","0.681","0.319","清晰"],
  ["厦门","0.562","0.864","0.954","1.000","0.438","清晰"],
  ["泉州","0.289","0.227","0.058","0.076","0.231","清晰"],
  ["宁德","0.009","0.000","0.049","0.226","0.226","清晰"],
  ["漳州","0.024","0.227","0.022","0.045","0.206","可辨"],
  ["莆田","0.163","0.000","0.049","0.038","0.163","可辨"],
  ["龙岩","0.120","0.046","0.111","0.102","0.074","弱（噪声）"],
  ["三明","0.096","0.046","0.079","0.040","0.056","弱（噪声）"],
  ["南平","0.034","0.046","0.055","0.026","0.029","弱（噪声）"]],
 [0.6,1.0,1.0,1.0,1.0,0.7,1.3], fontsize=9)

heading("5.2  障碍度方法与主障碍识别（含手算示范）", 2)
para("障碍度份额的含义是：某环节“离满分还差多少”，占该市“所有环节差距之和”的比例；份额最大的环节即“主障碍”。"
     "计算分三步（以 5.1 的标准化值 vₖ 为输入，理想值=1）：①偏离度 = 1−vₖ（离前沿差多少）；②总偏离 = Σ(1−vⱼ)"
     "（四环节差距相加）；③障碍度份额 = 偏离度 ÷ 总偏离。四环节份额必然加总为 100%。")
para("示范①——福州（教育侧封顶、产业侧偏低，主障碍清晰）：", before=4, after=2, bold=True, size=10)
worked_table("福州", [1.000,1.000,0.858,0.681])
para("福州教育两环节已达前沿（偏离 0 → 障碍 0%），失调全部来自产业侧，其中产业转化占总差距 69.2% → 主障碍=产业转化。",
     before=2, size=9.5)
para("示范②——龙岩（四环节同步偏低，份额趋同，主障碍为弱信号）：", before=4, after=2, bold=True, size=10)
worked_table("龙岩", [0.120,0.046,0.111,0.102])
para("龙岩四环节 vₖ 都很低（0.05–0.12）→ 偏离度都大且彼此接近（0.88–0.95）→ 除以总偏离后各≈25%。"
     "这正是腹地“各占 25%”的来历：不是巧合，而是四环节差得差不多（故其主障碍判为弱信号，见 5.3）。", before=2, size=9.5)
para("两点补充：①vₖ 来自该环节人均值在全样本（9市×6年）的 min-max 位置，0=最低、1=前沿；"
     "②学术式“障碍度=因子贡献度×指标偏离度”中，贡献度=权重×偏离度，本文四环节等权（各 1/4），"
     "归一化时 1/4 上下抵消，故份额 = (1−vₖ)/Σ(1−vⱼ)——等权下权重不影响份额。", before=4, size=9.5)
figure("/tmp/fig3_obstacle.png", 6.1, "图3  障碍度四环节构成（B口径，2024）。双核障碍集中单侧，腹地多市四环节近均分。")

heading("5.3  障碍度检验：min-max 压缩效应及其修正", 2)
para("检验动机：图3 中腹地多市四环节障碍度近似各占 25%，需检验其是否为真实“均衡薄弱”，抑或方法假象。")
p=doc.add_paragraph(); p.paragraph_format.line_spacing=1.4
setfont(p.add_run("机制诊断："), bold=True, size=10)
setfont(p.add_run("障碍度份额 = (1−vₖ)/Σ(1−vⱼ)。由于福州、厦门占据 min-max 的上界，全部腹地城市被压缩进 [0, 0.29] 的"
 "底部区间（见 5.1 表）。在该压缩区内，若一城四环节标准化值彼此接近（如南平 0.034/0.046/0.055/0.026，极差仅 0.029），"
 "则 (1−vₖ) 皆≈0.9 → 障碍份额机械地皆≈25%，其“主障碍”由第三位小数决定，本质为噪声。"), size=10)
p=doc.add_paragraph(); p.paragraph_format.line_spacing=1.4
setfont(p.add_run("判定准则（极差阈值）："), bold=True, size=10)
setfont(p.add_run("以四环节标准化值的极差作为障碍可辨性指标。极差>0.16（福州、厦门、泉州、宁德、漳州、莆田）的城市，"
 "存在可识别的主障碍，分解可信；极差<0.08（龙岩、三明、南平）的城市，四环节同步贴近地板、无显著单一卡点，"
 "其主障碍标注应判为弱信号。“各占 25%”仅对后 3 市成立——这本身是发现（四环节同步赤贫的发育不足），"
 "而非可分解的结构性偏科。"), size=10)
p=doc.add_paragraph(); p.paragraph_format.line_spacing=1.4
setfont(p.add_run("修正与稳健性："), bold=True, size=10)
setfont(p.add_run("①报告障碍度时并列极差/集中度，对极差<0.08 的城市标注“四环节均衡薄弱·无显著主卡点”，"
 "不强行指定主障碍（本文表与图已据此将龙岩/三明/南平标为弱信号）；②稳健性可对障碍度改用不压缩底部的标准化"
 "（秩标准化、对数化或对绝对理想值的偏离），复核腹地主障碍排序是否稳定。该检验表明：双核的障碍诊断稳健，"
 "腹地的障碍“方向”多数可辨（宁德、莆田卡供给质量等），但其“份额量值”受 min-max 压缩、不宜过度解读。"), size=10)
para("修正落地——双轨障碍诊断（数据处理）：", before=6, after=2, bold=True, size=10)
para("第①步，取四环节人均值（供给规模/产业承接/产业转化÷常住人口；供给质量为计数）。"
     "第②步（轨A·全样本）：在全 9 市×6 年共 54 个单元上对每个环节做 min-max，得 vₖᴬ，衡量"
     "“该市离全省前沿的差距”。第③步（轨B·腹地组内）：剔除福州、厦门，仅在腹地 7 市×6 年共 42 个单元上"
     "对每个环节重新 min-max，得 vₖᴮ，衡量“腹地城市彼此之间的相对短板”。第④步：两轨各自代入同一障碍度公式"
     "(1−vₖ)/Σ(1−vⱼ)，得各自主障碍。轨B 因剔除了双核极值，腹地四环节不再被压到地板，标准化值极差大幅放大。",
     size=9.5)
para("检验过程：逐市比较轨A 与轨B 的主障碍是否一致——一致则该市诊断稳健（两种基准都指向同一短板）；"
     "不一致则说明其主障碍受标准化基准影响，须分别表述。下表同时列两轨主障碍与极差（极差越大、主障碍越可辨）。",
     before=2, size=9.5)
table(["市","轨A 全样本主障碍","轨B 腹地组内主障碍","是否一致","极差A","极差B"],
 [["泉州","产业承接","产业承接","一致","0.231","0.694"],
  ["龙岩","供给质量","供给质量","一致","0.074","0.383"],
  ["三明","产业转化","产业转化","一致","0.056","0.242"],
  ["莆田","供给质量","供给质量","一致","0.163","0.296"],
  ["宁德","供给质量","供给质量","一致","0.226","0.966"],
  ["漳州","产业承接","供给规模","不一致","0.206","0.956"],
  ["南平","产业转化","供给规模","不一致","0.029","0.225"]],
 [0.7,1.9,2.1,1.0,0.85,0.85], fontsize=9)
para("7 市中 5 市两轨一致（诊断稳健）；漳州、南平两轨不一致——全样本主障碍反映“离全省前沿”的最大缺口，"
     "组内主障碍反映“腹地内部”相对短板，二者并存须分别表述，不可简单择一。", before=2, size=9.5)

heading("5.4  典型个案", 2)
for h,t in [("福厦镜像","双核同为高协调，障碍却分处两端：福州待产业（转化）承接其人才，厦门待教育（规模）扩张其供给。"),
 ("宁德","产业指数（0.137，宁德时代专利井喷拉动）远高于教育指数（0.005），主障碍为供给质量——产业先行、本地教育尚未接续。"),
 ("莆田","耦合度逐年下滑（0.338→0.244），与其招生规模 2021 年后收缩一致，主障碍为供给质量。"),
 ("泉州型口径漏洞","泉州 E>I 系“AI 数字产业”指标定义所致（其信息软件业从业、AI 专利人均仅为福厦的约 1/8）；"
  "但泉州制造业内嵌的 AI 人才吸纳（智能制造）计入门类 C、不在门类 I，故产业承接对制造业大市存在系统性低估，"
  "宜以“制造业数字化就业”作稳健性补充。")]:
    p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(3); p.paragraph_format.line_spacing=1.4
    setfont(p.add_run(f"{h}："), bold=True, size=10); setfont(p.add_run(t), size=10)

heading("5.5  权重稳健性：等权与熵值法对照", 2)
para("四环节是否等权？本文以等权为基准（四环节在功能链中地位等同，任一缺失即失调），并以熵值法作稳健性。", after=2)
para("熵值法权重的数据处理（按环节 j 逐列计算）：", before=2, after=2, bold=True, size=10)
para("第①步，取该环节在全样本 54 个单元的标准化值 vᵢ。第②步，归一为比重 pᵢ = vᵢ / Σᵢvᵢ。"
     "第③步，算信息熵 eⱼ = −(1/ln n)·Σᵢ pᵢ·ln pᵢ（n=54；约定 pᵢ=0 时该项为 0）——某环节城际差异越小，熵越大。"
     "第④步，算信息效用 dⱼ = 1−eⱼ（差异越大、效用越大）。第⑤步，归一为权重 wⱼ = dⱼ / Σⱼdⱼ。所得四环节熵权如下表。",
     size=9.5)
table(["环节","熵值法权重","等权"],
 [["供给规模","0.170","0.250"],["供给质量","0.296","0.250"],
  ["产业承接","0.246","0.250"],["产业转化","0.289","0.250"]],
 [2.2,1.6,1.6], fontsize=9.5)
para("以熵权重算 2024 年 D：双核（福州、厦门）在两种权重下均稳居前 2，9 市排名仅三明与漳州相邻互换、D 值几乎不变"
     "（福州 0.936→0.934、厦门 0.913→0.927）。结论：双核—腹地格局对权重选择稳健，等权法作主模型成立。"
     "熵值法下供给规模权重偏低、供给质量/产业转化偏高，源于福厦在后两者上离散度更大——故熵值法宜作稳健性而非唯一主模型。",
     before=2, size=9.5)

# ===== 六 空间扩展:置换检验 =====
heading("六、空间扩展：厦漳泉都市圈协同的三市组合置换检验", 1)
para("“双核—腹地”是按行政市的静态分化，但 AI 人才供需匹配未必止于行政边界。为检验厦门、漳州、泉州是否存在"
     "都市圈协同，且避免小样本空间计量的设定争议，本文采用穷举置换检验：在 9 市任选 3 市的全部 C(9,3)=84 个组合中，"
     "看厦漳泉的区域协同表现是否高于多数随机组合。该法零新数据、不依赖空间权重矩阵。")
para("数据处理与指标：对每个三市组合，将三市并为一个区域，区域 E、I 用人口加权聚合"
     "（E_组=ΣpᵢEᵢ/Σpᵢ；对人均环节恰等于“三市合并后人均”，且沿用 9 市原标准化基准，故 D_组 与单市 Dᵢ 同尺度可比）。"
     "据此构造三个指标：①区域协调度 D_组=√(C·T)；②区域协同溢价 RSP=D_组−mean(Dᵢ)（组团是否优于单市平均）；"
     "③均衡改善度 BI=mean(|Eᵢ−Iᵢ|)−|E_组−I_组|（组团是否使供需更均衡）——BI 为协同互补的主指标。", before=2)
para("检验过程：对 84 组逐一计算三指标并排序，看厦漳泉组合的名次与分位。结果如下表。", before=2, after=2)
table(["指标","含义","厦漳泉排名","解读"],
 [["D_组","区域整体协调度","49 / 84（前58%）","中游——漳、泉低 D 拖累整体水平"],
  ["RSP","区域协同溢价","53 / 84（前63%）","简单“组团提升”不突出"],
  ["BI","均衡改善度","5 / 84（前6%）","强信号——互补使供需更均衡"]],
 [1.0,2.4,1.9,2.7], fontsize=9.5)
figure("/tmp/fig4_permutation.png", 6.0, "图4  厦漳泉在 84 组三市组合中的均衡改善度（BI）排位（B口径，2024）。")
para("解读：厦漳泉在“整体水平（D_组）”与“简单协同溢价（RSP）”上仅居中游，但在“均衡改善度（BI）”上位列前 6%——"
     "即厦门产业领先、泉漳教育与承接相对偏强，三市并为区域后教育—产业缺口显著缩小，呈现互补型都市圈协同信号。", before=4)
para("结论边界（重要）：BI 高位组合普遍包含厦门（前 10% 组合中厦门出现 8/8），故这是“厦门锚定的互补效应”，"
     "厦漳泉因地理邻近与产业链联系而具备现实都市圈基础、落在高均衡区，但置换检验只能提示协同信号，"
     "不能证明因果“抱团”；其机制仍需跨市人才流动、企业合作与产业链数据进一步验证。", before=2, size=9.5)

# ===== 七 =====
heading("七、稳健性与局限", 1)
for t in ["三口径：C/X 口径作稳健性对照（C 口径一流专业稀疏，近纯规模驱动）。",
 "窗口上界 2024：2025 因年鉴未出版＋专利右截断不纳入；即便 2024，专利亦存轻微右截断，最末 1—2 年产业转化谨慎解读。",
 "回推单元：人口 2019、漳州 2024 产业承接为全省定锚回推，作稳健性起点；莆田 2019 产业承接代理值已剔除。",
 "标准化零值与压缩：min-max 最小值单元 E=0 触发“不可计算”（待定 ε 下限）；底部压缩使腹地障碍份额趋同（见 5.3）。",
 "供给质量：当前复合仅含一流专业成分（0.6），学位点（0.4）待补，补后双核教育指数将进一步抬升。"]:
    p=doc.add_paragraph(style=None); p.paragraph_format.left_indent=Inches(0.25); p.paragraph_format.space_after=Pt(2); p.paragraph_format.line_spacing=1.35
    setfont(p.add_run("· "), size=10); setfont(p.add_run(t), size=10)

para("（初稿 · 2026-06-22 · 数值取自管道实跑结果 latest_functional_chain_ccd_results.csv，随面板更新与口径冻结迭代。）",
     size=8.5, color=(0x88,0x88,0x88), before=10)

out="/Users/greenbarry/Documents/GitHub/fj-esystem-and-isystem-CCD/manuscripts/论文A_数据与检验.docx"
doc.save(out); print("saved", out)

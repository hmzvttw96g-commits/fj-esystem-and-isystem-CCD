# -*- coding: utf-8 -*-
"""论文A 初稿（CSSCI 体例）构建器。复用数据与检验的排版体例 + 实跑结果 + 5 图。
参考文献仅列已核实条目；需引而未核实处一律以【待补文献：…】标注，绝不杜撰。
用法：python3 scripts/build_paper_a_manuscript.py  → manuscripts/论文A初稿.docx
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pathlib import Path

ROOT = Path(__file__).resolve().parents[1]
FIG = ROOT / "manuscripts" / "figures"
BODY = "Songti SC"; HEAD = "Heiti SC"
doc = Document()
sec = doc.sections[0]
sec.page_width = Inches(8.27); sec.page_height = Inches(11.69)
for m in ("top_margin", "bottom_margin"): setattr(sec, m, Inches(1.0))
for m in ("left_margin", "right_margin"): setattr(sec, m, Inches(1.1))
_fp = sec.footer.paragraphs[0]; _fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
_fld = OxmlElement('w:fldSimple'); _fld.set(qn('w:instr'), ' PAGE ')
_r = OxmlElement('w:r'); _rpr = OxmlElement('w:rPr'); _sz = OxmlElement('w:sz'); _sz.set(qn('w:val'), '18'); _rpr.append(_sz)
_co = OxmlElement('w:color'); _co.set(qn('w:val'), '888888'); _rpr.append(_co); _r.append(_rpr)
_t = OxmlElement('w:t'); _t.text = '1'; _r.append(_t); _fld.append(_r); _fp._p.append(_fld)


def setfont(run, name=BODY, size=10.5, bold=False, color=None):
    run.font.name = name; run.font.size = Pt(size); run.bold = bold
    if color: run.font.color.rgb = RGBColor(*color)
    rpr = run._element.get_or_add_rPr(); rf = rpr.find(qn('w:rFonts'))
    if rf is None: rf = OxmlElement('w:rFonts'); rpr.append(rf)
    rf.set(qn('w:eastAsia'), name); rf.set(qn('w:ascii'), name); rf.set(qn('w:hAnsi'), name)


def para(text="", size=10.5, bold=False, align=None, before=2, after=4, color=None, name=BODY, indent=None):
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(before); p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.5
    p.alignment = align if align else WD_ALIGN_PARAGRAPH.JUSTIFY
    do_ind = indent if indent is not None else (align is None and not bold)
    if do_ind: p.paragraph_format.first_line_indent = Pt(size * 2)
    if text: setfont(p.add_run(text), name=name, size=size, bold=bold, color=color)
    return p


def runs(p, segs):
    """在段落 p 上追加多段 run：segs=[(text,bold,color), ...]"""
    for seg in segs:
        t = seg[0]; b = seg[1] if len(seg) > 1 else False; c = seg[2] if len(seg) > 2 else None
        setfont(p.add_run(t), size=10.5, bold=b, color=c)


def _accent_rule(p, color="378ADD", sz=10):
    pPr = p._p.get_or_add_pPr(); b = OxmlElement('w:pBdr'); bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'), 'single'); bot.set(qn('w:sz'), str(sz)); bot.set(qn('w:space'), '4'); bot.set(qn('w:color'), color)
    b.append(bot); pPr.append(b)


def heading(text, lvl=1):
    sizes = {1: 15, 2: 12.5, 3: 11.5}
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(14 if lvl == 1 else 8); p.paragraph_format.space_after = Pt(6 if lvl == 1 else 5)
    p.paragraph_format.keep_with_next = True
    setfont(p.add_run(text), name=HEAD, size=sizes[lvl], bold=True, color=(0x1F, 0x3A, 0x4A))
    if lvl == 1: _accent_rule(p)
    return p


def shade(cell, hexc):
    tcPr = cell._tc.get_or_add_tcPr(); sh = OxmlElement('w:shd')
    sh.set(qn('w:val'), 'clear'); sh.set(qn('w:fill'), hexc); tcPr.append(sh)


def table(headers, rows, widths, header_fill="1F3A4A", header_color=(255, 255, 255), fontsize=9.5, zebra=True):
    t = doc.add_table(rows=1, cols=len(headers)); t.alignment = WD_TABLE_ALIGNMENT.CENTER; t.style = 'Table Grid'
    for j, h in enumerate(headers):
        c = t.rows[0].cells[j]; c.width = Inches(widths[j]); shade(c, header_fill)
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        setfont(c.paragraphs[0].add_run(h), name=HEAD, size=fontsize, bold=True, color=header_color)
    for ri, row in enumerate(rows):
        cells = t.add_row().cells
        for j, val in enumerate(row):
            cells[j].width = Inches(widths[j])
            if zebra and ri % 2 == 1: shade(cells[j], "F2F5F7")
            pp = cells[j].paragraphs[0]; pp.alignment = WD_ALIGN_PARAGRAPH.CENTER if j > 0 else WD_ALIGN_PARAGRAPH.LEFT
            setfont(pp.add_run(str(val)), size=fontsize)
    return t


def figure(path, width, caption):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before = Pt(6)
    p.add_run().add_picture(str(path), width=Inches(width))
    cap = doc.add_paragraph(); cap.alignment = WD_ALIGN_PARAGRAPH.CENTER; cap.paragraph_format.space_after = Pt(8)
    setfont(cap.add_run(caption), name=BODY, size=9, color=(0x66, 0x66, 0x66))


def todo(text):
    """待补文献/待补内容的醒目标注（橙色），提醒作者补全，绝不杜撰。"""
    p = doc.add_paragraph(); p.paragraph_format.left_indent = Inches(0.25); p.paragraph_format.space_after = Pt(3); p.paragraph_format.line_spacing = 1.4
    setfont(p.add_run("【待补】"), size=9.5, bold=True, color=(0xC0, 0x5A, 0x1D))
    setfont(p.add_run(text), size=9.5, color=(0xC0, 0x5A, 0x1D))


# ============ 题名 / 作者 / 摘要 ============
ti = doc.add_paragraph(); ti.alignment = WD_ALIGN_PARAGRAPH.CENTER
setfont(ti.add_run("教育—产业耦合协调的功能链测度与障碍诊断"), name=HEAD, size=18, bold=True, color=(0x1F, 0x3A, 0x4A))
st = doc.add_paragraph(); st.alignment = WD_ALIGN_PARAGRAPH.CENTER; st.paragraph_format.space_after = Pt(4)
setfont(st.add_run("——以福建省人工智能人才供需匹配为例"), name=HEAD, size=13, bold=True, color=(0x37, 0x8A, 0xDD))
au = doc.add_paragraph(); au.alignment = WD_ALIGN_PARAGRAPH.CENTER; au.paragraph_format.space_after = Pt(2)
setfont(au.add_run("【作者姓名】　【作者单位，城市　邮编】"), size=10, color=(0x55, 0x55, 0x55))
fu = doc.add_paragraph(); fu.alignment = WD_ALIGN_PARAGRAPH.CENTER; fu.paragraph_format.space_after = Pt(10)
setfont(fu.add_run("【基金项目：项目名称及编号（待补）】"), size=9, color=(0x88, 0x88, 0x88))

p = doc.add_paragraph(); p.paragraph_format.line_spacing = 1.5; p.paragraph_format.space_after = Pt(4)
setfont(p.add_run("摘　要："), bold=True, size=10)
setfont(p.add_run(
    "人工智能人才的教育供给与产业需求能否在区域内匹配，是数字经济高质量发展的关键约束。本文将“教育—产业”关系"
    "解构为“供给规模—供给质量—产业承接—产业转化”四环节功能链，以耦合协调度（CCD）模型测度福建省 9 个设区市"
    "2019—2024 年的教育（E）—产业（I）耦合协调水平，并以障碍度模型诊断失调的关键环节。研究发现：①福建 AI 人才"
    "供需匹配呈显著的“双核—腹地”二元结构——福州、厦门达优质协调（D>0.9），其余 7 市长期处于失调区间；②福州“教育"
    "领先”、厦门“产业领先”构成结构镜像，二者跨核互补最强；③腹地城市的失调多为“功能链整体发育不足”而非单点错配；"
    "④边界约束的区域组合检验表明，区域均衡改善主要来自“教育领先↔产业领先”的跨极配对，纯腹地城市彼此组团几乎不"
    "产生协同，提示腹地需接入核心而非“抱团取暖”。本文的边际贡献在于：将整体 CCD 拓展为功能链局部 CCD 以识别"
    "链节断裂，并以边界约束的穷举置换检验为区域协同提供可证伪的探索性证据。"), size=10)
p = doc.add_paragraph(); p.paragraph_format.line_spacing = 1.5; p.paragraph_format.space_after = Pt(10)
setfont(p.add_run("关键词："), bold=True, size=10)
setfont(p.add_run("人工智能人才；教育—产业耦合；耦合协调度；障碍度诊断；功能链；福建省"), size=10)
todo("英文题名、Abstract、Keywords 待译补；作者、单位、基金、JEL/中图分类号待补。")

# ============ 一、引言 ============
heading("一、引言", 1)
para("数字经济的竞争，归根结底是人才的竞争。人工智能（AI）作为新一轮科技革命与产业变革的核心驱动力，对人才的"
     "数量与质量提出了双重要求；而人才的形成依赖教育系统的持续供给，人才的价值实现依赖产业系统的有效吸纳与转化。"
     "当教育供给与产业需求在结构、质量与空间上无法匹配时，便出现“有人无岗”或“有岗无人”的结构性错配，制约区域"
     "创新与产业升级。因此，刻画并诊断区域内“教育—产业”的匹配状态，是理解人才供需矛盾、制定差异化政策的前提。")
para("既有研究多从单侧（教育规模或产业需求）描述人才供需，较少将二者置于同一“系统协调”框架下进行可比测度，"
     "亦较少深入到“教育—产业”内部链节，回答“在哪一环节、以何种形式发生错配”。这导致政策建议常停留于“扩大供给”"
     "或“引进人才”的笼统层面，缺乏对失调结构的精准识别。", )
para("福建省为观察上述问题提供了典型样本：省内既有福州、厦门两个高教与产业“双核”，又有教育与产业基础相对薄弱的"
     "广大腹地城市，且拥有宁德时代等具有全国影响力的 AI 相关产业主体，区域内部供需结构差异显著。", )
todo("引言需补：①福建数字经济/AI 产业规模与人才缺口的权威数据来源（如省统计公报、工信部门或第三方人才报告）；"
     "②“人才供需结构性错配”命题的代表性文献 2—3 篇。")
para("基于此，本文将“教育—产业”关系解构为“供给规模—供给质量—产业承接—产业转化”四环节功能链，以耦合协调度模型"
     "测度福建 9 市 2019—2024 年的 E—I 耦合协调水平，并以障碍度模型定位失调环节。在方法上做两点拓展：其一，将整体"
     "耦合协调度拓展为功能链链节之间的“局部耦合协调度”，以识别整体协调可能掩盖的链节断裂；其二，以“边界约束的"
     "区域组合穷举置换检验”，为城市群/都市圈层面的空间协同提供可证伪的探索性证据。本文余下结构为：第二节文献综述"
     "与理论分析，第三节研究设计，第四节实证结果，第五节稳健性检验，第六节结论与政策建议。")

# ============ 二、文献综述与理论分析 ============
heading("二、文献综述与理论分析", 1)
heading("2.1　人才供需匹配与产教融合", 2)
para("人才供需匹配与产教融合是本文的现实落点。已有讨论普遍指出，教育供给与产业需求之间存在结构性、质量性与空间性"
     "错配，产教融合的推进仍受制于校企利益协调机制不健全等因素。但相关研究多为政策分析与案例描述，量化测度区域"
     "“教育—产业”匹配程度、并诊断其失调来源的实证工作仍相对不足。")
todo("2.1 需补：人才供需匹配测度（如供需匹配指数、技能错配 mismatch）2—3 篇；产教融合/教育—产业协同的实证 2—3 篇；"
     "AI/数字人才的就业效应或供需研究 1—2 篇（搜索可见“人工智能就业效应”等主题真实存在，但须经 CNKI/期刊核出完整出处后引用）。")
heading("2.2　耦合协调度模型及其应用", 2)
p = para("")
runs(p, [("耦合协调度（Coupling Coordination Degree, CCD）模型源于物理学耦合概念，廖重斌（1999）", False),
         ("[1]", False, (0x18, 0x5F, 0xA5)),
         ("以珠江三角洲为例，建立了环境与经济协调发展的定量评判及分类体系，奠定了国内 CCD 的基本范式；此后该模型被"
          "广泛用于两系统乃至多系统的协调测度。针对其在应用中的常见偏误，王淑佳等（2021）", False),
         ("[2]", False, (0x18, 0x5F, 0xA5)),
         ("系统梳理了国内耦合协调度模型的“四类误区”并给出修正建议，为规范使用提供了依据。近年亦有研究将 CCD 用于"
          "子系统的两两配对诊断——如 Zhang 等（2025）", False),
         ("[3]", False, (0x18, 0x5F, 0xA5)),
         ("对经济—社会—基础设施—生态四个韧性子系统构造六个二元 CCD 以识别其互配关系，为“局部耦合协调”提供了先例。", False)])
para("在失调来源的诊断上，障碍度模型常与 CCD 配合使用，通过各指标对协调水平的“偏离贡献”识别关键短板（瓶颈），"
     "已被广泛应用于乡村振兴、城市韧性与数字经济等领域的协调诊断。", )
todo("2.2 需补：①障碍度模型的方法学出处（其在中文文献中的奠基/规范引用，须核出完整出处）；②熵值法客观赋权的"
     "方法学出处；③CCD 在教育/科教领域应用的代表文献 1—2 篇。")
heading("2.3　文献述评与本文定位", 2)
para("综上，既有研究为本文提供了方法基础与现实背景，但仍存三点可拓展空间：其一，多停留于“系统对系统”的整体协调，"
     "较少深入功能链内部、识别具体断裂的链节；其二，对 min-max 标准化在区域极值悬殊时对腹地诊断的“压缩效应”"
     "缺乏自觉处理；其三，城市群/都市圈层面的“空间协同”多为定性判断，缺乏可证伪的统计参照。本文针对性地以"
     "功能链局部 CCD、双轨障碍诊断与边界约束置换检验回应上述三点。")
heading("2.4　理论框架：四环节功能链与（拟纳入的）金融调节", 2)
para("本文将“教育—产业”人才匹配视为一条功能链：教育系统先以“供给规模”培养人才存量，再以“供给质量”决定人才的"
     "层次；产业系统先以“产业承接”吸纳人才就业，再以“产业转化”将人才转化为创新产出。四环节首尾相连，任一环节"
     "的薄弱都会沿链条传导，形成区域人才供需的整体协调或局部断裂。教育（E）由“供给规模、供给质量”合成，产业（I）"
     "由“产业承接、产业转化”合成，二者的耦合协调度刻画区域 AI 人才供需的总体匹配水平。")
p = para("")
runs(p, [("需要说明的是，金融（F）作为撬动教育投入与产业投资的要素，理论上对 E—I 耦合具有调节作用。", False),
         ("本文为“测度与诊断”定位，实证部分聚焦 E—I 耦合与障碍诊断，金融调节作为后续因果识别工作（另文）的研究"
          "对象，此处仅在框架中标明、不作经验估计，以免超出本文数据所能支撑的结论范围。", False, (0x55, 0x55, 0x55))])

# ============ 三、研究设计 ============
heading("三、研究设计", 1)
heading("3.1　指标体系", 2)
para("依据四环节功能链构建指标体系（表 1）。考虑“AI 人才”口径的模糊性，本文设置严格嵌套的三口径（C⊂B⊂X）："
     "C 为 AI 直接命名专业（人工智能、智能科学与技术、数据科学与大数据技术），B 在 C 基础上加入计算机科学与技术、"
     "软件工程等基础专业群（主口径），X 进一步纳入电子信息、自动化、智能制造等交叉群；专利端按 IPC 前缀对应嵌套。"
     "除供给质量（项目点计数）外，各指标以全市常住人口人均化，使其度量“人均强度”而非城市规模。", before=2)
table(["子系统", "环节", "测度指标", "数据来源", "处理"],
      [["E 教育", "供给规模", "AI 相关专业本专科招生数", "各校招生章程/计划（逐校核）", "人均化"],
       ["E 教育", "供给质量", "国家级/省级一流本科专业建设点", "教育部/省教育厅公开名单（逐条核验）", "容量（非人均）"],
       ["I 产业", "产业承接", "信息传输、软件业城镇单位从业人员（门类 I）", "各市统计年鉴（官方核验）", "人均化"],
       ["I 产业", "产业转化", "企业 AI/数字技术专利申请量", "专利数据库（IPC 检索）", "人均化"]],
      [0.8, 0.9, 2.3, 2.1, 0.7], fontsize=8.8)
para("表 1　四环节功能链指标体系", before=0, after=6, size=8.8, color=(0x66, 0x66, 0x66), indent=False, align=WD_ALIGN_PARAGRAPH.CENTER)
heading("3.2　数据来源、时间口径与缺失处理", 2)
p = para("")
runs(p, [("样本为福建 9 设区市 2019—2024 年面板。供给质量的时间口径采“认定年起、存量持续”计法：某一流本科专业"
          "建设点自其获批（认定）年起在其后各年持续计入（每年只计其自身权重一次，不滚动累加），依据为一流专业"
          "“双万计划”（2019—2021 三批）", False),
         ("[7]", False, (0x18, 0x5F, 0xA5)),
         ("系“建设”属性、认定后全周期持续监测、仅对严重质量不达标者动态撤销", False),
         ("[6]", False, (0x18, 0x5F, 0xA5)),
         ("——故认定即确立可持续的质量信号，而非待验收方计。", False)])
para("数据处理遵循“缺失不填 0、不插值、不外推”的原则；仅在“全省锚可得 + 单一年份”的条件下作定锚回推（如人口"
     "2019 年口径断点、个别市个别年份产业承接的定锚回推），并将其纳入稳健性起点而非主结论依据。专利数据经去重、"
     "剔除高校与个人申请、按 IPC 分类与第一申请人地址归市后入账。完备性核查对“零命中”单元生成真零核查单逐一复核"
     "（如经核宁德师范学院 2019—2021 三批省级一流名单确无 AI/计算机类专业入选，系真零而非缺失）。", )
heading("3.3　测度方法", 2)
p = para("")
runs(p, [("（1）耦合协调度。教育与产业的耦合协调度为 D=√(C·T)，其中耦合度 C=2√(E·I)/(E+I) 测两系统的同步性、"
          "综合发展度 T=(E+I)/2 测整体水平；D∈[0,1]，越大越协调。该形式与廖重斌（1999）", False),
         ("[1]", False, (0x18, 0x5F, 0xA5)),
         ("及王淑佳等（2021）", False),
         ("[2]", False, (0x18, 0x5F, 0xA5)),
         ("厘定的规范公式一致，并通过其“四类误区”自检。E、I 各由其两环节等权合成；环节值先按 C/B/X 口径在全样本"
          "（9 市 × 6 年）上作 min-max 标准化。", False)])
para("（2）障碍度诊断。对四环节，以“偏离度”v_k 的归一化份额衡量其对失调的贡献：障碍度份额 = (1−v_k)/Σ(1−v_j)，"
     "份额最大者即主障碍环节。等权设定下 1/4 在归一化中相互抵消，不影响份额排序。", )
p = para("")
runs(p, [("（3）局部耦合协调度。为识别整体协调可能掩盖的链节断裂，本文对四环节两两构造“局部 CCD”（公式同上），"
          "主文取理论意义最强的四组：D₁₂（教育内部）、D₃₄（产业内部）、D₁₃（数量供需）、D₂₄（质量转化），其余两对"
          "（D₁₄、D₂₃）入附录以做满 C(4,2)=6 对，与 Zhang 等（2025）", False),
         ("[3]", False, (0x18, 0x5F, 0xA5)),
         ("的两两 CCD 思路一致。", False)])
p = para("")
runs(p, [("（4）边界约束的区域组合检验。区域协同不应在任意城市组合上比较，而须先确定“边界”（行政/规划/地理相邻），"
          "再在同规模随机组合（安慰剂参照）中看其分位——此即可变面积单元问题（MAUP）的要求", False),
         ("[4]", False, (0x18, 0x5F, 0xA5)),
         ("，空间连续性约束的区域化思想亦见诸 Assunção 等（2006）", False),
         ("[5]", False, (0x18, 0x5F, 0xA5)),
         ("。区域聚合采人口加权；定义协同溢价 RSP=D_组−mean(D_i)、均衡改善度 BI=mean|E_i−I_i|−|E_组−I_组|，"
          "以 BI 为协同互补主指标。", False)])

# ============ 四、实证结果 ============
heading("四、实证结果", 1)
heading("4.1　总体格局：“双核—腹地”二元结构", 2)
figure(FIG / "fig1_heatmap.png", 5.8, "图1　福建 9 市教育—产业耦合协调度热力图（B 口径，2019—2024）")
para("图 1 显示，福州、厦门 2024 年耦合协调度达 0.91 以上（优质协调），其余 7 市长期低于 0.4（失调区间），呈鲜明的"
     "“双核—腹地”二元结构。腹地城市中，泉州（0.36）相对最高，宁德（0.16）最低。从时序看，双核保持高位稳定，"
     "莆田等市则呈下滑（如莆田由 2020 年的 0.35 降至 2024 年的 0.26），与其招生规模收缩相关。", before=2)
heading("4.2　E—I 结构与双核镜像", 2)
figure(FIG / "fig2_scatter.png", 4.7, "图2　E-I 结构与双核镜像（B 口径，2024）")
para("将各市置于 E—I 平面（图 2）可见双核的结构性差异：福州偏“教育领先”（E>I），厦门偏“产业领先”（I>E），二者"
     "构成关于均衡线的镜像。腹地 7 市则聚集于原点附近的“双低”区域，教育与产业供给同步偏弱。这一镜像结构是后文"
     "“跨核互补”最强的微观基础。", before=2)
heading("4.3　障碍度诊断与压缩效应", 2)
figure(FIG / "fig3_obstacle.png", 5.8, "图3　障碍度四环节构成（B 口径，2024）")
para("障碍度分解（图 3）显示双核的相对短板各异（福州在产业转化、厦门在教育内部相关环节），与其镜像结构呼应。"
     "但须警惕方法假象：由于福州、厦门占据 min-max 标准化的上界，腹地城市被压缩进区间底部，其四环节偏离度"
     "皆趋近 1、障碍份额机械地皆趋近 25%。因此本文以四环节标准化值的“极差”作为障碍可辨性指标——极差<0.08 的"
     "城市（龙岩、三明、南平）四环节同步贴近地板、无显著单一卡点，其“主障碍”应判为弱信号，“各占 25%”本身即是"
     "“四环节同步发育不足”的发现，而非可分解的结构性偏科。", before=2)
heading("4.4　功能链局部断裂", 2)
figure(FIG / "fig5_localccd.png", 5.7, "图5　整体 CCD 与四组功能链局部 CCD（B 口径，2024）")
para("局部 CCD（图 5）进一步揭示整体协调所掩盖的链节结构：双核“全链协调”、无明显断裂；泉州呈“局部优势·总体"
     "低位”，其教育内部 D₁₂（0.51）明显高于其他链节，即“教育端内部尚协调、但产业接不住”，是可培育的突破口；"
     "龙岩、三明、南平四组局部 CCD 同步低位，印证其失调系功能链整体发育不足。需要特别区分两类“零值断裂”：宁德"
     "供给质量为真零（经核确无 AI 类一流），其含质量链节 D₁₂、D₂₄=0 系真实的质量缺失；而莆田原亦为 0，经补入"
     "其 2020 年获批的计算机省级一流后已非零（D₁₂=0.29、D₂₄=0.20），属此前的数据缺口而非结构断裂——提示"
     "“质量链节为 0”须先甄别“真零/缺失”再作解读。", before=2)
heading("4.5　空间协同的边界约束检验", 2)
figure(FIG / "fig4_permutation.png", 5.7, "图4　厦漳泉在 84 组三市组合（同规模安慰剂分布）中的均衡改善度（BI）排位（B 口径，2024）")
para("在 C(9,3)=84 个三市组合的安慰剂分布中（图 4），现实边界组合呈三种协同类型：福州—厦门“跨核互补”在两市"
     "组合中 D_S、BI 均列第 1（结构互补最强，但二者非地理相邻，定位为理论互补对照）；厦漳泉“互补型都市圈”BI 居"
     "前 7%（厦门产业领先 + 泉漳教育/承接相对偏强，并区后供需缺口缩小）；福州都市圈、闽东北“带动型”RSP 居同规模"
     "前 2—5%（由福州高水平拉动，但 BI 中低、互补不强）。", before=2)
para("安慰剂分布本身亦给出一般规律：含核心市（福/厦）组合的 D_组均值（0.696）是纯腹地组合（0.292）的约 2.4 倍，"
     "纯腹地组合 BI 均值≈0（彼此组团几乎不产生均衡改善，最高 D 组仅 0.352、仍失调）；跨极（同时含教育领先与产业"
     "领先城市）组合 BI 均值为正、不跨极为负。这为“腹地弱联动”提供反事实证据：均衡改善来自“跨越教育领先↔产业"
     "领先”的镜像配对，腹地需接入核心而非“抱团取暖”。", )
para("结论边界（重要）：以上为相对同规模随机参照的探索性信号，不证明因果协同；区域协同的因果机制仍需跨市人才"
     "流动、企业合作与产业链数据进一步验证。", color=(0x55, 0x55, 0x55), size=10)

# ============ 五、稳健性检验 ============
heading("五、稳健性检验", 1)
p = para("")
runs(p, [("本文从四方面检验主结论的稳健性（详细数值见《数据与检验》技术附件）。①双轨障碍诊断：在剔除福厦、于腹地"
          "7 市组内重新标准化后重算障碍度，腹地内部相对短板的方向与全样本基本一致，缓解了压缩效应对腹地诊断的干扰；"
          "②客观赋权：以熵值法替代等权重算 2024 年 D，双核稳居前 2、9 市排名仅中段小幅重排；③耦合度形式：以 ε 下限"
          "（[0.01,1]，缓解零值地板）与平方式耦合度 4EI/(E+I)²（张等 2025）", False),
         ("[3]", False, (0x18, 0x5F, 0xA5)),
         ("替代主模型重算，双核在三种形式下均居前 2、与主模型的 Spearman 秩相关达 0.983；④边界置换：现实边界组合的"
          "协同结论在同规模安慰剂分布中稳健。综合可见，“双核—腹地”格局对权重与耦合度形式选择稳健。本文方法亦通过"
          "王淑佳等（2021）", False),
         ("[2]", False, (0x18, 0x5F, 0xA5)),
         ("提出的“四类误区”自检。", False)])

# ============ 六、结论与政策建议 ============
heading("六、结论与政策建议", 1)
para("本文以四环节功能链测度并诊断了福建省 AI 人才的教育—产业耦合协调。主要结论：①福建 AI 人才供需匹配呈"
     "“双核—腹地”二元结构，双核优质协调、腹地长期失调；②福州“教育领先”、厦门“产业领先”构成镜像，跨核互补最强；"
     "③腹地失调多为功能链整体发育不足而非单点错配，且彼此组团难以自我协同；④区域均衡改善源于“教育领先↔产业领先”"
     "的跨极配对。", before=2)
para("据此提出差异化政策建议：其一，双核应强化“跨核互补”，推动福州的教育优势与厦门的产业优势对接（如联合培养、"
     "产业飞地），而非各自重复布局；其二，腹地城市不宜“抱团取暖”，应优先接入核心都市圈、承接其溢出，同时针对其"
     "“整体发育不足”采取“基础设施 + 基础专业 + 基础产业”的整链培育；其三，对供给质量为真零的城市（如宁德），"
     "应以引进或共建一流专业建设点补质量短板，避免“有规模无质量”的扩张型供给。", )
para("本文的局限与展望：①金融调节未纳入实证，后续将以教育政策冲击构建准实验、识别金融对 E—I 耦合的因果调节"
     "（另文）；②专利数据存在轻微右截断，最末 1—2 年的产业转化宜谨慎解读；③min-max 标准化的零值地板与压缩效应"
     "已以双轨与多形式稳健性缓解，但仍是测度的内在限制；④边界组合检验为探索性，因果协同有待流动数据验证；"
     "⑤供给质量复合目前仅含一流专业成分，学位点等成分待补后将进一步细化教育质量刻画。", )

# ============ 参考文献 ============
heading("参考文献", 1)
para("（仅列已核实条目；正文中以【待补】标注处，待经 CNKI/期刊原文核出完整出处后补入，本稿不以未核实文献充数。）",
     size=9, color=(0x88, 0x88, 0x88), after=4)
for ref in [
    "[1] 廖重斌. 环境与经济协调发展的定量评判及其分类体系——以珠江三角洲城市群为例[J]. 热带地理, 1999, 19(2): 171-177.",
    "[2] 王淑佳, 孔伟, 任亮, 等. 国内耦合协调度模型的误区及修正[J]. 自然资源学报, 2021, 36(3): 793-810.",
    "[3] Zhang M, Hu Y, Mao Y, et al. 经济—社会—基础设施—生态韧性系统的耦合协调研究：以浙江省为例[J]. PLOS ONE, 2025. doi:10.1371/journal.pone.0323673.（题名以原文为准）",
    "[4] Openshaw S. The Modifiable Areal Unit Problem[M]. Norwich: GeoBooks, 1984.（Concepts and Techniques in Modern Geography, No.38）",
    "[5] Assunção R M, Neves M C, Câmara G, et al. Efficient regionalization techniques for socio-economic geographical units using minimum spanning trees[J]. International Journal of Geographical Information Science, 2006, 20(7): 797-811.",
    "[6] 教育部, 财政部, 国家发展改革委. “双一流”建设成效评价办法（试行）[Z]. 2021.",
    "[7] 教育部办公厅. 关于实施一流本科专业建设“双万计划”的通知（教高厅函〔2019〕18号）[Z]. 2019."]:
    p = doc.add_paragraph(); p.paragraph_format.left_indent = Inches(0.3); p.paragraph_format.first_line_indent = Inches(-0.3)
    p.paragraph_format.space_after = Pt(3); p.paragraph_format.line_spacing = 1.3
    setfont(p.add_run(ref), size=9.5)
todo("尚需补充的文献类别（务必核出完整出处后再引、按上文【待补】各处对应补入）："
     "①人才供需匹配/技能错配测度；②产教融合/教育—产业协同实证；③AI/数字人才就业效应；④障碍度模型方法学出处；"
     "⑤熵值法赋权方法学出处；⑥福建数字经济/AI 产业与人才缺口的权威数据来源；⑦CCD 在教育/科教领域的应用文献。")

para("（论文初稿 · 2026-06-25 · 数值取自管道实跑结果，与《数据与检验》技术附件一致；参考文献仅列已核实条目，"
     "未核实处以【待补】标注，绝不杜撰。）", size=8.5, color=(0x88, 0x88, 0x88), before=10)

out = ROOT / "manuscripts" / "论文A初稿.docx"
doc.save(str(out))
print(f"saved {out}")

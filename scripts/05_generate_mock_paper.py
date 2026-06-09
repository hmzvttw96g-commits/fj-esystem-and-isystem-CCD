from __future__ import annotations

from pathlib import Path
from typing import Dict, List

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

from common import GENERATED_DIR, LOGS_DIR, PROCESSED_DIR, create_pdf_text, read_xlsx_rows, setup_logger, write_xlsx


logger = setup_logger("generate_mock_paper", LOGS_DIR / "generate_mock_paper.log")

FULL_SCORE = 100
STRUCTURE = [
    ("选择题", "一、选择题（本大题共15小题，每小题3分，共45分）"),
    ("填空题", "二、填空题（本大题共5小题，每小题3分，共15分）"),
    ("解答题", "三、解答题（本大题共4小题，每小题10分，共40分）"),
]


QUESTIONS: List[Dict[str, object]] = [
    {
        "id": "1",
        "type": "选择题",
        "score": 3,
        "kp": "集合与逻辑",
        "detail": "集合交并补运算",
        "diff": "基础",
        "text": "设集合 A={1,3,5,7}，B={3,4,5,6}，则 A∩B=（  ）。\nA. {1,7}  B. {3,5}  C. {1,3,4,5,6,7}  D. {4,6}",
        "answer": "B",
        "solution": "交集取两个集合共有元素，A与B共有3和5，所以A∩B={3,5}。",
    },
    {
        "id": "2",
        "type": "选择题",
        "score": 3,
        "kp": "集合与逻辑",
        "detail": "区间与描述法",
        "diff": "基础",
        "text": "不等式 x>=2 且 x<6 的解集可表示为（  ）。\nA. (2,6)  B. [2,6)  C. (2,6]  D. [2,6]",
        "answer": "B",
        "solution": "x>=2 左端点取到，x<6 右端点取不到，所以为[2,6)。",
    },
    {
        "id": "3",
        "type": "选择题",
        "score": 3,
        "kp": "不等式",
        "detail": "一元二次不等式",
        "diff": "中等",
        "text": "不等式 (x-1)(x-4)<0 的解集是（  ）。\nA. x<1  B. x>4  C. 1<x<4  D. x<1或x>4",
        "answer": "C",
        "solution": "二次式开口向上，两根为1和4，小于0时取两根之间，故1<x<4。",
    },
    {
        "id": "4",
        "type": "选择题",
        "score": 3,
        "kp": "函数",
        "detail": "函数定义域",
        "diff": "中等",
        "text": "函数 y=sqrt(x-3) 的定义域是（  ）。\nA. x>=3  B. x>3  C. x<=3  D. 全体实数",
        "answer": "A",
        "solution": "根号内必须非负，x-3>=0，所以x>=3。",
    },
    {
        "id": "5",
        "type": "选择题",
        "score": 3,
        "kp": "函数",
        "detail": "函数求值",
        "diff": "基础",
        "text": "已知 f(x)=2x^2-1，则 f(2)=（  ）。\nA. 3  B. 5  C. 7  D. 9",
        "answer": "C",
        "solution": "f(2)=2×2^2-1=8-1=7。",
    },
    {
        "id": "6",
        "type": "选择题",
        "score": 3,
        "kp": "指数与对数",
        "detail": "指数式与对数式互化",
        "diff": "中等",
        "text": "将 3^2=9 化为对数式，正确的是（  ）。\nA. log_3 9=2  B. log_9 3=2  C. log_2 9=3  D. log_3 2=9",
        "answer": "A",
        "solution": "a^b=N 等价于 log_a N=b，所以3^2=9等价于log_3 9=2。",
    },
    {
        "id": "7",
        "type": "选择题",
        "score": 3,
        "kp": "三角函数",
        "detail": "特殊角三角函数",
        "diff": "基础",
        "text": "sin30°+cos60° 的值为（  ）。\nA. 0  B. 1/2  C. 1  D. 2",
        "answer": "C",
        "solution": "sin30°=1/2，cos60°=1/2，和为1。",
    },
    {
        "id": "8",
        "type": "选择题",
        "score": 3,
        "kp": "数列",
        "detail": "等差数列",
        "diff": "基础",
        "text": "等差数列 4,7,10,... 的公差是（  ）。\nA. 2  B. 3  C. 4  D. 7",
        "answer": "B",
        "solution": "相邻两项差为7-4=3。",
    },
    {
        "id": "9",
        "type": "选择题",
        "score": 3,
        "kp": "概率统计",
        "detail": "抽样与频率分布",
        "diff": "基础",
        "text": "某校从高一、高二、高三学生中按人数比例抽取样本，这种抽样方法通常称为（  ）。\nA. 简单随机抽样  B. 分层抽样  C. 系统抽样  D. 任意抽样",
        "answer": "B",
        "solution": "按不同层的人数比例抽样，属于分层抽样。",
    },
    {
        "id": "10",
        "type": "选择题",
        "score": 3,
        "kp": "立体几何",
        "detail": "立体几何视图与位置",
        "diff": "中等",
        "text": "一个正方体共有（  ）个面。\nA. 4  B. 6  C. 8  D. 12",
        "answer": "B",
        "solution": "正方体有6个面、8个顶点、12条棱。",
    },
    {
        "id": "11",
        "type": "选择题",
        "score": 3,
        "kp": "集合与逻辑",
        "detail": "充分必要条件",
        "diff": "中等",
        "text": "“x=2”是“x^2=4”的（  ）。\nA. 充分不必要条件  B. 必要不充分条件  C. 充要条件  D. 既不充分也不必要条件",
        "answer": "A",
        "solution": "x=2一定推出x^2=4；但x^2=4还可能x=-2，所以不是必要条件。",
    },
    {
        "id": "12",
        "type": "选择题",
        "score": 3,
        "kp": "直线与圆",
        "detail": "直线斜率",
        "diff": "基础",
        "text": "直线 y=-3x+5 的斜率是（  ）。\nA. -3  B. 3  C. 5  D. -5",
        "answer": "A",
        "solution": "直线y=kx+b中k为斜率，所以斜率为-3。",
    },
    {
        "id": "13",
        "type": "选择题",
        "score": 3,
        "kp": "向量",
        "detail": "向量坐标运算",
        "diff": "中等",
        "text": "已知向量 a=(1,2)，b=(3,-1)，则 a+b=（  ）。\nA. (4,1)  B. (2,3)  C. (3,2)  D. (-2,3)",
        "answer": "A",
        "solution": "向量加法按坐标相加，a+b=(1+3,2-1)=(4,1)。",
    },
    {
        "id": "14",
        "type": "选择题",
        "score": 3,
        "kp": "直线与圆",
        "detail": "圆的标准方程",
        "diff": "中等",
        "text": "圆 (x+1)^2+(y-2)^2=16 的半径是（  ）。\nA. 1  B. 2  C. 4  D. 16",
        "answer": "C",
        "solution": "标准方程右端为r^2=16，所以r=4。",
    },
    {
        "id": "15",
        "type": "选择题",
        "score": 3,
        "kp": "概率统计",
        "detail": "计数原理",
        "diff": "基础",
        "text": "从4种饮料和3种点心中各选1种，共有（  ）种不同搭配。\nA. 7  B. 12  C. 24  D. 43",
        "answer": "B",
        "solution": "分步计数：4×3=12种。",
    },
    {
        "id": "16",
        "type": "填空题",
        "score": 3,
        "kp": "不等式",
        "detail": "一元一次不等式",
        "diff": "基础",
        "text": "不等式 2x-5>=1 的解集为__________。",
        "answer": "x>=3",
        "solution": "2x-5>=1，2x>=6，所以x>=3。",
    },
    {
        "id": "17",
        "type": "填空题",
        "score": 3,
        "kp": "函数",
        "detail": "函数奇偶性",
        "diff": "中等",
        "text": "函数 f(x)=x^2+1 是__________函数。（填“奇”或“偶”）",
        "answer": "偶",
        "solution": "f(-x)=(-x)^2+1=x^2+1=f(x)，所以是偶函数。",
    },
    {
        "id": "18",
        "type": "填空题",
        "score": 3,
        "kp": "数列",
        "detail": "等比数列",
        "diff": "中等",
        "text": "等比数列首项为3，公比为2，则第4项为__________。",
        "answer": "24",
        "solution": "a4=3×2^(4-1)=24。",
    },
    {
        "id": "19",
        "type": "填空题",
        "score": 3,
        "kp": "概率统计",
        "detail": "概率基础",
        "diff": "基础",
        "text": "袋中有2个红球、3个白球，任取1个球，取到红球的概率为__________。",
        "answer": "2/5",
        "solution": "总数5个，其中红球2个，概率为2/5。",
    },
    {
        "id": "20",
        "type": "填空题",
        "score": 3,
        "kp": "直线与圆",
        "detail": "直线与圆综合",
        "diff": "中等",
        "text": "圆 x^2+y^2=25 与 x 轴正半轴的交点坐标为__________。",
        "answer": "(5,0)",
        "solution": "x轴上y=0，代入得x^2=25。正半轴取x=5，所以交点为(5,0)。",
    },
    {
        "id": "21",
        "type": "解答题",
        "score": 10,
        "kp": "集合与逻辑",
        "detail": "集合交并补运算",
        "diff": "中等",
        "text": "已知全集 U={1,2,3,4,5,6,7,8}，集合 A={2,4,6,8}，B={1,2,3,4}。\n(1) 求 A∩B；\n(2) 求 A∪B；\n(3) 求集合 A 在全集 U 中的补集。",
        "answer": "(1){2,4}；(2){1,2,3,4,6,8}；(3){1,3,5,7}",
        "solution": "A∩B为共有元素{2,4}。A∪B为两个集合所有元素组成的集合{1,2,3,4,6,8}。U中不属于A的元素为{1,3,5,7}。",
    },
    {
        "id": "22",
        "type": "解答题",
        "score": 10,
        "kp": "函数",
        "detail": "分段函数",
        "diff": "较难",
        "text": "已知分段函数 f(x)={ x+2, x<1；2x-1, x>=1 }。\n(1) 求 f(0) 和 f(3)；\n(2) 若 f(x)=5，求x；\n(3) 判断点(1,1)是否在该函数图像上，并说明理由。",
        "answer": "(1)f(0)=2，f(3)=5；(2)x=3；(3)在图像上",
        "solution": "0<1，用f(0)=0+2=2；3>=1，用f(3)=2×3-1=5。若x<1，则x+2=5得x=3不符合；若x>=1，则2x-1=5得x=3符合。x=1时f(1)=2×1-1=1，所以点(1,1)在图像上。",
    },
    {
        "id": "23",
        "type": "解答题",
        "score": 10,
        "kp": "直线与圆",
        "detail": "直线与圆综合",
        "diff": "较难",
        "text": "已知直线 l 经过点 A(1,2)、B(3,6)，圆 C 的方程为 (x-2)^2+(y-1)^2=9。\n(1) 求直线 l 的斜率；\n(2) 写出圆 C 的圆心和半径；\n(3) 判断点 B 是否在圆 C 上。",
        "answer": "(1)2；(2)圆心(2,1)，半径3；(3)不在圆上",
        "solution": "斜率k=(6-2)/(3-1)=2。圆的标准方程可知圆心为(2,1)，半径为3。点B代入：(3-2)^2+(6-1)^2=1+25=26，不等于9，所以不在圆上。",
    },
    {
        "id": "24",
        "type": "解答题",
        "score": 10,
        "kp": "应用与建模",
        "detail": "应用函数建模",
        "diff": "较难",
        "text": "某班准备购买学习资料。每本资料进价8元，另需固定配送费60元。设购买x本资料的总费用为y元。\n(1) 写出y关于x的函数关系式；\n(2) 购买35本需要多少元；\n(3) 若总预算不超过500元，最多可购买多少本？",
        "answer": "(1)y=8x+60；(2)340元；(3)55本",
        "solution": "总费用=单本费用×本数+固定配送费，所以y=8x+60。x=35时，y=8×35+60=340。8x+60<=500，8x<=440，x<=55，因此最多55本。",
    },
]


def load_ols_weight_lookup() -> Dict[str, Dict[str, object]]:
    path = PROCESSED_DIR / "detailed_knowledge_ols_analysis.xlsx"
    if not path.exists():
        return {}
    return {str(r.get("detailed_knowledge_point")): r for r in read_xlsx_rows(path, "ols_weighted_blueprint")}


def load_trend_lookup() -> Dict[str, Dict[str, object]]:
    path = PROCESSED_DIR / "future_exam_trend_prediction.xlsx"
    if not path.exists():
        return {}
    return {str(r.get("knowledge_point")): r for r in read_xlsx_rows(path, "trend_prediction")}


def add_title(doc: Document, title: str, subtitle: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(18)
    r.font.color.rgb = RGBColor(31, 78, 121)
    sp = doc.add_paragraph()
    sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sp.add_run(subtitle)
    sr.font.size = Pt(10)


def build_docx(path: Path, answers: bool = False) -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)
    doc.styles["Normal"].font.size = Pt(10.5)
    add_title(
        doc,
        "福建省中职学业水平测试公共基础数学模拟卷" + (" 答案解析" if answers else ""),
        f"按最新指定结构生成｜满分{FULL_SCORE}分｜15选择+5填空+4解答",
    )
    if not answers:
        doc.add_paragraph(
            f"说明：本卷只包含公共基础综合卷中的数学部分，为原创模拟卷。满分{FULL_SCORE}分，建议用时约70分钟。"
        )
    for qtype, intro in STRUCTURE:
        doc.add_heading(intro, level=1)
        for q in [x for x in QUESTIONS if x["type"] == qtype]:
            if answers:
                p = doc.add_paragraph()
                p.add_run(f"{q['id']}. 答案：{q['answer']}").bold = True
                doc.add_paragraph(f"解析：{q['solution']}")
            else:
                doc.add_paragraph(f"{q['id']}. {q['text']}")
    doc.save(path)


def pdf_sections(answers: bool = False):
    sections = []
    for qtype, _intro in STRUCTURE:
        paras = []
        for q in [x for x in QUESTIONS if x["type"] == qtype]:
            paras.append(f"{q['id']}. 答案：{q['answer']}。解析：{q['solution']}" if answers else f"{q['id']}. {q['text']}")
        sections.append({"heading": qtype, "paragraphs": paras})
    return sections


def validate_structure() -> None:
    counts = {
        "选择题": sum(1 for q in QUESTIONS if q["type"] == "选择题"),
        "填空题": sum(1 for q in QUESTIONS if q["type"] == "填空题"),
        "解答题": sum(1 for q in QUESTIONS if q["type"] == "解答题"),
    }
    score = sum(int(q["score"]) for q in QUESTIONS)
    expected = {"选择题": 15, "填空题": 5, "解答题": 4}
    if counts != expected or score != FULL_SCORE:
        raise ValueError(f"模拟卷结构错误：counts={counts}, score={score}")


def main() -> None:
    logger.info("Start generating OLS-weighted public basic math mock paper")
    validate_structure()
    GENERATED_DIR.mkdir(parents=True, exist_ok=True)
    paper_docx = GENERATED_DIR / "福建省中职学业水平测试公共基础数学模拟卷.docx"
    answer_docx = GENERATED_DIR / "福建省中职学业水平测试公共基础数学模拟卷_答案解析.docx"
    build_docx(paper_docx, answers=False)
    build_docx(answer_docx, answers=True)
    subtitle = f"按最新指定结构生成｜满分{FULL_SCORE}分｜15选择+5填空+4解答"
    create_pdf_text(GENERATED_DIR / "福建省中职学业水平测试公共基础数学模拟卷.pdf", "福建省中职学业水平测试公共基础数学模拟卷", pdf_sections(False), subtitle)
    create_pdf_text(GENERATED_DIR / "福建省中职学业水平测试公共基础数学模拟卷_答案解析.pdf", "福建省中职学业水平测试公共基础数学模拟卷 答案解析", pdf_sections(True), subtitle)

    trend_lookup = load_trend_lookup()
    ols_lookup = load_ols_weight_lookup()
    basis = []
    for q in QUESTIONS:
        trend = trend_lookup.get(str(q["kp"]), {})
        ols = ols_lookup.get(str(q["detail"]), {})
        weight = ols.get("normalized_weight", "")
        recommended = ols.get("recommended_questions_in_24", "")
        basis.append(
            {
                "question_id": q["id"],
                "question_type": q["type"],
                "score": q["score"],
                "knowledge_point": q["kp"],
                "detailed_knowledge_point": q["detail"],
                "difficulty": q["diff"],
                "ols_normalized_weight": weight,
                "ols_recommended_questions_in_24": recommended,
                "historical_real_share": ols.get("historical_real_share", ""),
                "xiamen_reference_share": ols.get("xiamen_reference_share", ""),
                "prediction_score": trend.get("prediction_score", ""),
                "recommended_coverage": trend.get("recommended_coverage", ""),
                "design_basis": "按真题逐题详细知识点与同年厦门卷重合/相关分析建立OLS权重；再结合考纲覆盖、历史真题频次和2026厦门模拟卷信号配置题目。题面为原创，题型结构严格按15选、5填、4解生成。",
                "answer": q["answer"],
            }
        )
    write_xlsx({"blueprint": basis}, GENERATED_DIR / "模拟卷命题依据表.xlsx")
    logger.info("Output generated paper, answer and OLS blueprint")


if __name__ == "__main__":
    main()
